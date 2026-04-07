from http.server import BaseHTTPRequestHandler
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys
from datetime import datetime
import re

# Add the api directory to path so we can import template_config
sys.path.insert(0, os.path.dirname(__file__))

try:
    from template_config import TEMPLATE_URLS
    print("[WILL] Successfully imported TEMPLATE_URLS")
    print(f"[WILL] Will template URL: {TEMPLATE_URLS.get('will', 'NOT FOUND')}")
except ImportError as e:
    print(f"[WILL] CRITICAL: Failed to import template_config: {e}")
    TEMPLATE_URLS = {'will': 'ERROR_NO_CONFIG'}

# Module-level cache to avoid re-downloading templates on warm invocations
_template_cache = {}

def download_template(url):
    """Download template from Google Drive, with in-memory caching.

    Uses requests for reliable redirect and cookie handling. Google Drive
    sometimes returns an HTML confirmation page for larger files; this
    function detects that and retries with the embedded confirm token.
    """
    if url in _template_cache:
        print(f"[WILL] Using cached template for: {url}")
        return BytesIO(_template_cache[url])
    try:
        import requests as req_lib
        print(f"[WILL] Downloading template from: {url}")
        session = req_lib.Session()
        session.headers['User-Agent'] = 'Mozilla/5.0'
        response = session.get(url, allow_redirects=True, timeout=30)
        response.raise_for_status()
        content = response.content
        if content[:4] != b'PK\x03\x04':
            confirm = re.search(rb'confirm=([0-9A-Za-z_\-]+)', content)
            file_id = re.search(r'[?&]id=([^&]+)', url)
            if confirm and file_id:
                retry_url = (
                    f'https://drive.google.com/uc?export=download'
                    f'&id={file_id.group(1)}'
                    f'&confirm={confirm.group(1).decode()}'
                )
                response = session.get(retry_url, allow_redirects=True, timeout=30)
                content = response.content
        if content[:4] != b'PK\x03\x04':
            raise Exception("Downloaded file is not a valid .docx (failed ZIP header check). "
                            "Ensure the file is shared as 'Anyone with the link can view'.")
        print(f"[WILL] Template downloaded: {len(content)} bytes")
        _template_cache[url] = content
        return BytesIO(content)
    except Exception as e:
        print(f"[WILL] Template download failed: {e}")
        raise Exception(f"Failed to download template: {str(e)}")

def _format_body_para(para):
    """Apply body-text formatting matching the Will template:
    Charter 12pt, justified, 1.5x line spacing, 1-inch first-line indent,
    6pt space before and after.
    """
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = Emu(914400)  # 1 inch — matches template body paragraphs
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    for run in para.runs:
        run.font.name = 'Charter'
        run.font.size = Pt(12)


def _format_heading_para(para):
    """Apply article-heading formatting matching the Will template:
    Charter 14pt bold, centered, 1.0x line spacing, 6pt space before and after.
    """
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing = 1.0
    pf.first_line_indent = Emu(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    for run in para.runs:
        run.font.name = 'Charter'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.small_caps = True


def format_name_for_filename(full_name):
    """Format name as 'Lastname Firstname' removing middle initials"""
    name_parts = full_name.strip().split()
    if len(name_parts) < 2:
        return full_name

    # Remove middle initials (single letter followed by optional period)
    filtered_parts = [part for part in name_parts if not re.match(r'^[A-Z]\.?$', part)]

    if len(filtered_parts) >= 2:
        # Last name is last element, first name is everything before it
        lastname = filtered_parts[-1]
        firstname = ' '.join(filtered_parts[:-1])
        return f"{lastname} {firstname}"
    else:
        return full_name

def int_to_roman(num):
    """Convert integer to Roman numeral"""
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syms = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syms[i]
            num -= val[i]
        i += 1
    return roman_num

def renumber_articles(doc, start_from_para_index, starting_article_num):
    """Renumber all articles starting from given paragraph index"""
    article_pattern = re.compile(r'^Article\s+([IVXLCDM]+)\s+-\s+(.+)$')
    current_article_num = starting_article_num

    for i in range(start_from_para_index, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        match = article_pattern.match(para.text)
        if match:
            article_title = match.group(2)
            new_roman = int_to_roman(current_article_num)
            para.text = f'Article {new_roman} - {article_title}'
            current_article_num += 1

def add_page_numbers(doc):
    """Add page numbers to footer"""
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = 1  # Center
        footer_para.text = "Page "
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

def load_clause_text(clause_filename):
    """Load clause text from text file"""
    clause_path = os.path.join(os.path.dirname(__file__), 'clauses', clause_filename)
    try:
        with open(clause_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except Exception as e:
        print(f"Error loading clause {clause_filename}: {e}")
        return None

def replace_in_runs(para, old_text, new_text):
    """Replace text while preserving formatting"""
    if old_text not in para.text:
        return
    
    # Get full text
    full_text = para.text
    new_full_text = full_text.replace(old_text, new_text)
    
    # Find all runs and their text
    runs = para.runs
    if not runs:
        return
    
    # Clear all run text
    for run in runs:
        run.text = ''
    
    # Put all text in first run to preserve its formatting
    if runs:
        runs[0].text = new_full_text

def replace_in_document(doc, replacements):
    """Replace all placeholders in document while preserving formatting"""
    for para in doc.paragraphs:
        para_text = para.text
        for key, value in replacements.items():
            if key in para_text:
                replace_in_runs(para, key, str(value))
    
    # Also replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = para.text
                    for key, value in replacements.items():
                        if key in para_text:
                            replace_in_runs(para, key, str(value))

def handle_conditional_blocks(doc, data, children=None):
    """Handle ##IF_MARRIED## and similar conditional blocks"""
    is_married = data.get('IS_MARRIED', True)
    has_contingent_beneficiary = bool(data.get('CONTINGENT_BENEFICIARY_NAME', '').strip())

    # Check if trust for minors exists (any child under 25)
    has_trust_for_minors = False
    if children:
        for child in children:
            age = calculate_age(child.get('dob', ''))
            if age < 25:
                has_trust_for_minors = True
                break

    paragraphs_to_remove = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text

        # Handle ##Delete first sentence if unmarried##
        if '##Delete first sentence if unmarried##' in text and not is_married:
            # Remove everything before this marker
            sentences = text.split('.')
            if len(sentences) > 1:
                # Keep everything after the marker
                new_text = text.split('##Delete first sentence if unmarried##')[1]
                para.text = new_text.strip()

        # Handle ##If no contingent beneficiary, replace with:##
        if '##If no contingent beneficiary, replace with:' in text:
            if not has_contingent_beneficiary:
                # Extract the replacement text
                match = re.search(r'##If no contingent beneficiary, replace with:\s*([^#]+)##', text)
                if match:
                    replacement = match.group(1).strip()
                    # Replace entire paragraph with the fallback text
                    para.text = text.split('##If no contingent beneficiary')[0].strip() + ' ' + replacement
            else:
                # Remove the conditional marker
                para.text = re.sub(r'\s*##If no contingent beneficiary[^#]+##', '', text)

        # Handle ##If trust for minors exists:##
        if '##If trust for minors exists:' in text:
            if has_trust_for_minors:
                # Extract the trust reference text
                match = re.search(r'##If trust for minors exists:\s*([^#]+)##', text)
                if match:
                    trust_ref = match.group(1).strip()
                    # Remove the entire marker first
                    clean_text = re.sub(r'\s*##If trust for minors exists:[^#]+##', '', text)
                    # Then append the trust reference after "per stirpes"
                    # Make sure we add the period back if it was removed
                    if 'per stirpes.' in clean_text:
                        para.text = clean_text.replace('per stirpes.', f'per stirpes, {trust_ref}')
                    elif 'per stirpes' in clean_text:
                        para.text = clean_text.replace('per stirpes', f'per stirpes, {trust_ref}')
            else:
                # Remove the conditional marker only
                para.text = re.sub(r'\s*##If trust for minors exists:[^#]+##', '', text)

        # Handle ##IF_MARRIED## blocks
        if '##IF_MARRIED##' in text and not is_married:
            paragraphs_to_remove.append(para)

        # Clean up any remaining ## markers
        if '##' in text:
            # Remove all ## markers
            cleaned = re.sub(r'##[^#]+##', '', text)
            if cleaned.strip():
                para.text = cleaned
            elif not cleaned.strip() and '##INSERT' not in text:
                paragraphs_to_remove.append(para)

    # Remove marked paragraphs
    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

def insert_article_iii_clauses(doc, data):
    """Insert optional clauses at ##INSERT_ARTICLE_III_CLAUSES## marker"""
    # Find the insertion point
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if '##INSERT_ARTICLE_III_CLAUSES##' in para.text:
            insert_index = i
            # Remove the marker paragraph
            p = para._element
            p.getparent().remove(p)
            break
    
    if insert_index is None:
        return
    
    # Build list of clauses to insert
    clauses_to_insert = []
    
    # 1. Love and Affection (Disinheritance) - FIRST
    if data.get('INCLUDE_DISINHERITANCE'):
        clause_text = load_clause_text('LWT_-_Clause_-_Love_and_Affection.txt')
        if clause_text:
            clause_text = clause_text.replace('{DISINHERITED_RELATION}', data.get('DISINHERITED_RELATION', ''))
            clause_text = clause_text.replace('{DISINHERITED_NAME}', data.get('DISINHERITED_NAME', ''))
            clauses_to_insert.append(clause_text)
    
    # 2. Handwritten List - SECOND
    if data.get('INCLUDE_HANDWRITTEN_LIST'):
        clause_text = load_clause_text('LWT_-_Clause_-_Handwritten_List.txt')
        if clause_text:
            clauses_to_insert.append(clause_text)
    
    # 3. Real Estate Debt - THIRD
    if data.get('INCLUDE_REAL_ESTATE_DEBT'):
        clause_text = load_clause_text('LWT_-_Clause_-_Real_Estate_Indebtedness.txt')
        if clause_text:
            clauses_to_insert.append(clause_text)

    # Note: No Contest is now inserted as a separate article (Article IV)
    # See insert_no_contest_article() function

    # Insert clauses
    for clause_text in clauses_to_insert:
        new_para = doc.add_paragraph(clause_text)
        _format_body_para(new_para)

        # Move paragraph to correct position
        new_para._element.getparent().remove(new_para._element)
        doc.paragraphs[insert_index]._element.addprevious(new_para._element)
        insert_index += 1

def insert_no_contest_article(doc, data):
    """Insert no-contest clause as Article IV if requested"""
    if not data.get('INCLUDE_NO_CONTEST'):
        return False

    # Find the insertion point (before Executor article)
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if '##INSERT_NO_CONTEST_ARTICLE##' in para.text:
            insert_index = i
            # Remove the marker
            p = para._element
            p.getparent().remove(p)
            break

    if insert_index is None:
        return False

    # Load no-contest clause text
    clause_text = load_clause_text('LWT_-_Clause_-_No_Contest_Provision.txt')
    if not clause_text:
        return False

    # Add Article IV header
    article_para = doc.add_paragraph('Article IV - No Contest Provision')
    _format_heading_para(article_para)

    article_para._element.getparent().remove(article_para._element)
    doc.paragraphs[insert_index]._element.addprevious(article_para._element)
    insert_index += 1

    # Add clause text
    clause_para = doc.add_paragraph(clause_text)
    _format_body_para(clause_para)

    clause_para._element.getparent().remove(clause_para._element)
    doc.paragraphs[insert_index]._element.addprevious(clause_para._element)
    insert_index += 1

    return insert_index  # Return index for renumbering

def insert_guardian_article(doc, data, minor_children, insert_index):
    """Insert guardian appointment article for children under 18.

    minor_children is a pre-filtered list of children whose age < 18.
    Returns updated insert_index.
    """
    if not minor_children:
        return insert_index

    num_minors = len(minor_children)
    child_word = 'child' if num_minors == 1 else 'children'
    is_are = 'is' if num_minors == 1 else 'are'

    # Build names list for minor children only
    names_list = [c.get('name', '') for c in minor_children if c.get('name')]
    if len(names_list) == 1:
        names = names_list[0]
    elif len(names_list) == 2:
        names = f"{names_list[0]} and {names_list[1]}"
    else:
        names = ', '.join(names_list[:-1]) + f', and {names_list[-1]}'

    spouse_type = 'husband' if data.get('SPOUSE_GENDER', 'Female') == 'Male' else 'wife'
    g1_name = data.get('GUARDIAN_NAME_1', '')
    g1_rel  = data.get('GUARDIAN_RELATION_1', '')
    g2_name = data.get('GUARDIAN_NAME_2', '')
    g2_rel  = data.get('GUARDIAN_RELATION_2', '')
    g3_name = data.get('GUARDIAN_NAME_3', '')
    g3_rel  = data.get('GUARDIAN_RELATION_3', '')

    guardian_text = (
        f"In the event my {spouse_type} predeceases me and my {child_word} {names} "
        f"{is_are} less than Eighteen (18) years of age, I direct that my {g1_rel}, "
        f"{g1_name}, serve as Guardian until such beneficiary reaches the age of majority. "
        f"Should {g1_name} be unwilling or unable to serve, I appoint my {g2_rel}, "
        f"{g2_name}, as Guardian until such beneficiary reaches the age of majority. "
        f"Should {g2_name} be unwilling or unable to serve, I appoint my {g3_rel}, "
        f"{g3_name}, as Guardian until such beneficiary reaches the age of majority."
    )

    article_para = doc.add_paragraph('Article IV - Appointment of Guardian')
    _format_heading_para(article_para)
    article_para._element.getparent().remove(article_para._element)
    doc.paragraphs[insert_index]._element.addprevious(article_para._element)
    insert_index += 1

    body_para = doc.add_paragraph(guardian_text)
    _format_body_para(body_para)
    body_para._element.getparent().remove(body_para._element)
    doc.paragraphs[insert_index]._element.addprevious(body_para._element)
    insert_index += 1

    return insert_index


def insert_trust_for_minors(doc, data, children, insert_index):
    """Insert trust for minors article at insert_index if any child is under 25.

    The ##INSERT_NEW_ARTICLES## marker has already been removed by the caller.
    Returns updated insert_index (or original index if no trust needed).
    """
    has_minor = any(calculate_age(c.get('dob', '')) < 25 for c in children)
    if not has_minor:
        return insert_index

    trust_text = load_clause_text('LWT_-_Trust_for_Minor_Children.txt')
    if not trust_text:
        return insert_index

    trust_text = trust_text.replace('{TRUSTEE_NAME}', data.get('TRUSTEE_NAME', ''))

    article_para = doc.add_paragraph('Article VI - Trust for Minor Children')
    _format_heading_para(article_para)
    article_para._element.getparent().remove(article_para._element)
    doc.paragraphs[insert_index]._element.addprevious(article_para._element)
    insert_index += 1

    for paragraph_text in trust_text.split('\n\n'):
        if paragraph_text.strip():
            new_para = doc.add_paragraph(paragraph_text.strip())
            _format_body_para(new_para)
            new_para._element.getparent().remove(new_para._element)
            doc.paragraphs[insert_index]._element.addprevious(new_para._element)
            insert_index += 1

    return insert_index

def calculate_age(dob_string):
    """Calculate age from birthdate string.

    Accepts both YYYY-MM-DD (native date input value) and Month DD, YYYY
    (the formatted string the form used to send before the fix).
    """
    if not dob_string:
        return 0
    for fmt in ('%Y-%m-%d', '%B %d, %Y', '%b %d, %Y'):
        try:
            dob = datetime.strptime(dob_string.strip(), fmt)
            today = datetime.now()
            return today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
        except (ValueError, TypeError):
            continue
    return 0

def format_children_list(children):
    """Format children for the will with proper readability"""
    if not children:
        return '', '', 'no'

    detailed_list = []
    simple_list = []

    for child in children:
        name = child.get('name', '')
        dob = child.get('dob', '')

        if name:
            simple_list.append(name)
            if dob:
                # Format date nicely
                try:
                    date_obj = datetime.strptime(dob, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%B %d, %Y')
                    detailed_list.append(f"{name}, born {formatted_date}")
                except (ValueError, TypeError) as e:
                    # If date parsing fails, just use the name
                    detailed_list.append(name)
            else:
                detailed_list.append(name)

    # Format detailed list with semicolons for better readability in legal documents
    # Example: "John Doe, born January 15, 2000; Jane Doe, born March 20, 2002; and Michael Doe, born June 10, 2005"
    if len(detailed_list) == 0:
        detailed = ''
    elif len(detailed_list) == 1:
        detailed = detailed_list[0]
    elif len(detailed_list) == 2:
        detailed = f"{detailed_list[0]} and {detailed_list[1]}"
    else:
        # Multiple children: use semicolons to separate entries for clarity
        detailed = '; '.join(detailed_list[:-1]) + f'; and {detailed_list[-1]}'

    # Format simple list with commas and "and" for the last item
    if len(simple_list) == 0:
        simple = ''
    elif len(simple_list) == 1:
        simple = simple_list[0]
    elif len(simple_list) == 2:
        simple = f"{simple_list[0]} and {simple_list[1]}"
    else:
        simple = ', '.join(simple_list[:-1]) + f', and {simple_list[-1]}'

    # Convert number to word
    num_words = ['zero', 'one', 'two', 'three', 'four', 'five', 'six', 'seven', 'eight', 'nine', 'ten']
    if len(children) <= 10:
        num_word = num_words[len(children)]
    else:
        num_word = str(len(children))

    return detailed, simple, num_word

def generate_will_document(data):
    """Generate Last Will and Testament using the template"""

    # Download template from Google Drive
    template_url = TEMPLATE_URLS.get('will', '')
    if not template_url or template_url == 'ERROR_NO_CONFIG':
        return {'error': 'Will template URL not configured'}

    try:
        template_buffer = download_template(template_url)
        doc = Document(template_buffer)
    except Exception as e:
        return {'error': f'Could not load template: {str(e)}'}
    
    # Parse children data
    children = []
    if data.get('children'):
        children = json.loads(data['children']) if isinstance(data['children'], str) else data['children']
    
    # Format children
    children_detailed, children_simple, num_children_word = format_children_list(children)
    
    # Determine pronouns and titles
    is_male = data.get('CLIENT_GENDER', 'Male') == 'Male'
    spouse_is_male = data.get('SPOUSE_GENDER', 'Female') == 'Male'
    
    # Build replacements dictionary with the template's variable names
    replacements = {
        '{CLIENT_NAME}': data.get('CLIENT_NAME', ''),
        '{CLIENT_COUNTY}': data.get('COUNTY', ''),
        '{COUNTY}': data.get('COUNTY', ''),
        '{SPOUSE_NAME}': data.get('SPOUSE_NAME', ''),
        '{NUMBER_OF_CHILDREN}': str(len(children)),
        '{CHILDREN_LIST}': children_simple,
        '{SPOUSE_TYPE}': 'husband' if spouse_is_male else 'wife',
        '{CLIENT_PRONOUN_SUBJECTIVE}': 'he' if is_male else 'she',
        '{CLIENT_PRONOUN_POSSESSIVE}': 'his' if is_male else 'her',
        '{he/she}': 'he' if is_male else 'she',
        '{his/her}': 'his' if is_male else 'her',
        '{SPOUSE_PRONOUN}': 'he' if spouse_is_male else 'she',
        '{SPOUSE_PRONOUN_POSSESSIVE}': 'his' if spouse_is_male else 'her',
        '{TESTATOR_TITLE}': 'Testator' if is_male else 'Testatrix',
        '{EXECUTOR_TITLE}': 'Executor' if is_male else 'Executrix',
        '{PRIMARY_EXECUTOR}': data.get('PRIMARY_EXECUTOR', ''),
        '{ALTERNATE_EXECUTOR}': data.get('ALTERNATE_EXECUTOR', ''),
        '{ALTERNATE_EXECUTOR_NAME}': data.get('ALTERNATE_EXECUTOR', ''),
        '{ALTERNATE_EXECUTOR_RELATION}': data.get('ALTERNATE_EXECUTOR_RELATION', ''),
        '{ALTERNATE_EXECUTOR_COUNTY}': data.get('ALTERNATE_EXECUTOR_COUNTY', ''),
        '{ALTERNATE_EXECUTOR_STATE}': data.get('ALTERNATE_EXECUTOR_STATE', 'Tennessee'),
        '{CONTINGENT_BENEFICIARY_NAME}': data.get('CONTINGENT_BENEFICIARY_NAME', ''),
        '{CONTINGENT_BENEFICIARY_RELATION}': data.get('CONTINGENT_BENEFICIARY_RELATION', ''),
        '{EXEC_MONTH}': data.get('EXECUTION_MONTH', ''),
        '{EXEC_YEAR}': data.get('EXECUTION_YEAR', ''),
        '{EXECUTION_MONTH}': data.get('EXECUTION_MONTH', ''),
        '{EXECUTION_YEAR}': data.get('EXECUTION_YEAR', ''),
    }
    
    # For the family status line, use detailed list
    replacements['{NUM_CHILDREN}'] = f"{num_children_word} ({len(children)})" if children else "no"

    # If children exist, build the detailed description
    if children:
        replacements['{CHILDREN_DESCRIPTION}'] = children_detailed
        replacements['{CHILDREN_DETAILED}'] = children_detailed
    else:
        replacements['{CHILDREN_DESCRIPTION}'] = ''
        replacements['{CHILDREN_DETAILED}'] = ''
    
    # Step 1: Handle conditional blocks
    handle_conditional_blocks(doc, data, children)

    # Step 2: Replace all variables
    replace_in_document(doc, replacements)
    
    # Step 2b: Insert Specific Bequests paragraphs before "A. To My Spouse".
    # Placeholders use {braces} format for Curly (https://curly.io), the Word
    # add-in used by this firm to fill in one-off variables in generated documents.
    # Looks for ##INSERT_SPECIFIC_BEQUESTS## marker first; falls back to finding
    # the "A. To My Spouse" paragraph directly (for Drive templates without the marker).
    specific_bequests_paragraphs = []
    if data.get('INCLUDE_SPECIFIC_BEQUESTS'):
        count = int(data.get('SPECIFIC_BEQUEST_COUNT') or 2)
        count = max(1, min(count, 26))
        letters = 'abcdefghijklmnopqrstuvwxyz'
        specific_bequests_paragraphs.append('I make the following specific bequests:')
        for i in range(count):
            specific_bequests_paragraphs.append(f'({letters[i]}) {{Specific_bequest_{letters[i]}}}')
        specific_bequests_paragraphs.append(
            'If any of the beneficiaries named above shall predecease me or does not '
            'survive me by thirty (30) days, the bequest to them shall lapse and become '
            'a part of my Residuary Estate.'
        )

    def _insert_bequests_before(anchor_elem):
        """Insert all specific bequest paragraphs immediately before anchor_elem.
        Repeated addprevious(anchor) calls produce correct order: each new element
        slides in just before the anchor, after the previously inserted one.
        """
        for para_text in specific_bequests_paragraphs:
            new_para = doc.add_paragraph(para_text)
            _format_body_para(new_para)
            new_para._element.getparent().remove(new_para._element)
            anchor_elem.addprevious(new_para._element)

    # Try marker first
    marker_found = False
    for para in doc.paragraphs:
        if '##INSERT_SPECIFIC_BEQUESTS##' in para.text:
            marker_elem = para._element
            if specific_bequests_paragraphs:
                _insert_bequests_before(marker_elem)
            marker_elem.getparent().remove(marker_elem)
            marker_found = True
            break

    # Fallback: if no marker, find "A. To My Spouse" and insert before it
    if not marker_found and specific_bequests_paragraphs:
        for para in doc.paragraphs:
            if para.text.strip().startswith('A.') and 'Spouse' in para.text:
                _insert_bequests_before(para._element)
                break

    # Step 3: Insert Article III clauses
    insert_article_iii_clauses(doc, data)

    # Step 3b: Handle ##INSERT_EXECUTOR_EXTRA## marker — replace with Sell Real
    # Estate clause if checked, otherwise just remove the marker.
    sell_clause_text = load_clause_text('LWT_-_Clause_-_Sell_Real_Estate.txt') if data.get('INCLUDE_SELL_REAL_ESTATE') else None
    for para in doc.paragraphs:
        if '##INSERT_EXECUTOR_EXTRA##' in para.text:
            marker_elem = para._element
            if sell_clause_text:
                new_para = doc.add_paragraph(sell_clause_text)
                _format_body_para(new_para)
                new_para._element.getparent().remove(new_para._element)
                marker_elem.addprevious(new_para._element)
            marker_elem.getparent().remove(marker_elem)
            break

    # Step 4: Insert no-contest as Article IV if requested
    no_contest_inserted = insert_no_contest_article(doc, data)

    # Step 5: Insert optional articles (guardian + trust) at ##INSERT_NEW_ARTICLES## marker.
    # Find and remove the marker once, then insert articles in document order:
    # guardian (children < 18) first, trust (children < 25) second.
    if children:
        articles_index = None
        for i, para in enumerate(doc.paragraphs):
            if '##INSERT_NEW_ARTICLES##' in para.text:
                articles_index = i
                para._element.getparent().remove(para._element)
                break
        if articles_index is not None:
            minor_children = [c for c in children if calculate_age(c.get('dob', '')) < 18]
            articles_index = insert_guardian_article(doc, data, minor_children, articles_index)
            articles_index = insert_trust_for_minors(doc, data, children, articles_index)

    # Step 6: Renumber articles from IV onwards based on what was inserted.
    # Articles I-III are fixed; everything from IV onwards gets sequential numbering.
    # NOTE: setting para.text wipes run formatting, so we re-apply heading formatting
    # to every article paragraph in Step 6b below.
    article_pattern = re.compile(r'^Article\s+([IVXLCDM]+)\s+-\s+(.+)$')
    article_num = 4  # Start from IV
    found_article_iv = False

    for para in doc.paragraphs:
        match = article_pattern.match(para.text.strip())
        if match and found_article_iv:
            article_title = match.group(2)
            new_roman = int_to_roman(article_num)
            para.text = f'Article {new_roman} - {article_title}'
            article_num += 1
        elif match:
            current_num_text = match.group(1)
            article_title = match.group(2)
            if current_num_text in ['IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']:
                found_article_iv = True
                new_roman = int_to_roman(article_num)
                para.text = f'Article {new_roman} - {article_title}'
                article_num += 1

    # Step 6b: Re-apply heading formatting to ALL article headings.
    # The renumbering step above calls para.text = '...' which wipes run formatting.
    # Articles IV+ from both the template and inserted clauses need this pass.
    for para in doc.paragraphs:
        if article_pattern.match(para.text.strip()):
            _format_heading_para(para)

    # Step 7: Add page numbers if not already there
    add_page_numbers(doc)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            # Generate document
            result = generate_will_document(data)
            
            if isinstance(result, dict) and 'error' in result:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(json.dumps(result).encode())
                return
            
            # Format filename as: YYYY-MM-DD LWT lastname firstname.docx
            today = datetime.now().strftime('%Y-%m-%d')
            formatted_name = format_name_for_filename(data.get("CLIENT_NAME", "Document"))
            filename = f"{today} LWT {formatted_name}.docx"

            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()

            self.wfile.write(result.getvalue())
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            error_response = json.dumps({'error': str(e)})
            self.wfile.write(error_response.encode())
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
