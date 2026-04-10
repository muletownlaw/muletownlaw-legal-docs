# api/probate_utils.py
"""Shared utilities for probate document generation."""
import os
import zipfile
from io import BytesIO
from datetime import datetime, timedelta
from docx import Document


# --- Pronoun & Title Derivation ---

def derive_pronouns(gender):
    """Return subject/possessive/object pronouns from gender."""
    if gender == 'Male':
        return {'subject': 'he', 'possessive': 'his', 'object': 'him'}
    return {'subject': 'she', 'possessive': 'her', 'object': 'her'}


def derive_pr_title(estate_type, pr_gender):
    """Return PR title: Executor/Executrix/Administrator/Administratrix."""
    titles = {
        ('Testate', 'Male'): 'Executor',
        ('Testate', 'Female'): 'Executrix',
        ('Intestate', 'Male'): 'Administrator',
        ('Intestate', 'Female'): 'Administratrix',
    }
    try:
        return titles[(estate_type, pr_gender)]
    except KeyError:
        raise ValueError(
            f"Unknown estate_type/pr_gender combination: "
            f"({estate_type!r}, {pr_gender!r}). "
            f"estate_type must be 'Testate' or 'Intestate'; "
            f"pr_gender must be 'Male' or 'Female'."
        )


# --- Date Formatting ---

def format_date_legal(date_str):
    """Convert YYYY-MM-DD to 'January 1, 2026' format."""
    dt = datetime.strptime(date_str, '%Y-%m-%d')
    return f"{dt.strftime('%B')} {dt.day}, {dt.strftime('%Y')}"


def ordinal_day(date_str):
    """Return ordinal day: '10th', '1st', '2nd', '3rd', etc."""
    dt = datetime.strptime(date_str, '%Y-%m-%d')
    day = dt.day
    if 11 <= day <= 13:
        suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix}"


# --- Merge Field Replacement ---

def merge_runs_in_paragraph(paragraph):
    """Merge all runs to handle Word's split placeholders."""
    if not paragraph.runs:
        return
    full_text = paragraph.text
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = full_text


def replace_in_document(doc, replacements):
    """Replace all placeholders in a document. Merges runs first."""
    for paragraph in doc.paragraphs:
        merge_runs_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    merge_runs_in_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)


def load_template(template_name):
    """Load a .docx template from api/probate-templates/."""
    template_dir = os.path.join(os.path.dirname(__file__), 'probate-templates')
    path = os.path.join(template_dir, template_name)
    return Document(path)


# --- ZIP Assembly ---

def build_zip(documents, date_str=None):
    """Build a ZIP file from a list of (filename, Document) tuples.

    Args:
        documents: list of (filename_without_date, docx.Document) tuples
        date_str: YYYY-MM-DD string for filename prefix. Defaults to today.

    Returns:
        BytesIO buffer containing the ZIP file.
    """
    if date_str is None:
        date_str = datetime.now().strftime('%Y-%m-%d')

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for filename, doc in documents:
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            full_name = f"{date_str} {filename}.docx"
            zf.writestr(full_name, doc_buffer.getvalue())
    zip_buffer.seek(0)
    return zip_buffer


# --- Document Selection Logic (spec § 2.1) ---

def select_opening_documents(data):
    """Return list of (template_filename, output_title) for opening docs.

    Implements spec § 2.1 document selection logic.
    """
    estate_type = data['estate_type']
    docs = []

    if data.get('small_estate_election'):
        docs.append(('Small Estate - Affidavit as to Small Estate CURLY.docx',
                      'Affidavit as to Small Estate'))
        docs.append(('Small Estate - Order Approving Small Estate CURLY.docx',
                      'Order Approving Small Estate'))
        return docs

    if estate_type == 'Testate':
        if data.get('muniment_only'):
            docs.append(('Petition for Muniment of Title.docx',
                          'Petition for Muniment of Title'))
            docs.append(('Order for Muniment of Title.docx',
                          'Order for Muniment of Title'))
        elif data.get('will_type') == 'Standard Witnessed':
            docs.append(('Petition to Probate Will Ltrs Testamentary CURLY (1).docx',
                          'Petition to Probate Will and Letters Testamentary'))
            docs.append(('Order to Probate LWT CURLY.docx',
                          'Order to Probate Last Will and Testament'))
        elif data.get('will_type') == 'Holographic':
            docs.append(('Petition to Probate Holographic Will.docx',
                          'Petition to Probate Holographic Will'))
            docs.append(('Order Admitting Holographic LWT.docx',
                          'Order Admitting Holographic Last Will and Testament'))
        elif data.get('will_type') == 'Will + Codicil':
            docs.append(('Petition to Probate Will and Codicil.docx',
                          'Petition to Probate Will and Codicil'))
            docs.append(('Order Admitting Codicil and LWT.docx',
                          'Order Admitting Codicil and Last Will and Testament'))
        else:
            raise ValueError(
                f"Unknown will_type {data.get('will_type')!r} for Testate estate. "
                f"Expected 'Standard Witnessed', 'Holographic', or 'Will + Codicil'."
            )
    else:  # Intestate
        docs.append(('Petition for Appointment of Administrator CURLY.docx',
                      'Petition for Appointment of Administrator'))
        docs.append(('Order for Intestate Administration.docx',
                      'Order for Intestate Administration'))

    # Always add oath (for non-small-estate)
    docs.append(('Personal Representative Oath CURLY.docx',
                  'Personal Representative Oath'))

    return docs


# --- Declination Logic (spec § 2.2) ---

def determine_declinations(data):
    """Return list of dicts for people who need to sign declinations.

    Each dict: {name, relationship, gender, relation_to_pr}
    Implements spec § 2.2.
    """
    estate_type = data['estate_type']
    heirs = data.get('heirs', [])
    declinations = []

    if estate_type == 'Testate':
        if data.get('named_executor_can_serve', True):
            return []

        # Named executor can't serve — they must decline
        if data.get('will_executor_name'):
            declinations.append({
                'name': data['will_executor_name'],
                'relationship': 'Named Executor',
                'gender': data.get('will_executor_gender', 'Male'),
            })

        if data.get('will_names_alternate'):
            # Alternate exists and presumably can serve — no further declinations
            return declinations

        # No one in will can serve — need consent from each adult devisee/legatee
        for heir in heirs:
            if heir.get('heir_is_beneficiary') and not heir.get('heir_is_minor'):
                if heir['heir_full_name'] != data.get('pr_full_name'):
                    declinations.append({
                        'name': heir['heir_full_name'],
                        'relationship': heir['heir_relationship'],
                        'gender': heir.get('heir_gender', 'Male'),
                    })

    else:  # Intestate
        pr_relationship = data.get('pr_relationship', '')

        if pr_relationship == 'Spouse':
            return []  # Highest priority — no declinations

        surviving_spouse = data.get('decedent_spouse_name') if data.get(
            'decedent_marital_status') == 'Married' else None

        # For any non-spouse PR, the surviving spouse must decline
        if surviving_spouse and pr_relationship != 'Spouse':
            declinations.append({
                'name': surviving_spouse,
                'relationship': 'Surviving Spouse',
                'gender': 'Female' if data.get('decedent_gender') == 'Male' else 'Male',
            })

        if pr_relationship == 'Child':
            # Siblings (other children) must decline
            for heir in heirs:
                if (heir['heir_relationship'] in ['Son', 'Daughter', 'Child']
                        and heir['heir_full_name'] != data.get('pr_full_name')
                        and not heir.get('heir_is_minor')):
                    declinations.append({
                        'name': heir['heir_full_name'],
                        'relationship': heir['heir_relationship'],
                        'gender': heir.get('heir_gender', 'Male'),
                    })
        elif pr_relationship in ['Grandchild', 'Sibling', 'Other Kin', 'Unrelated']:
            # All children of decedent must decline
            for heir in heirs:
                if (heir['heir_relationship'] in ['Son', 'Daughter', 'Child']
                        and heir['heir_full_name'] != data.get('pr_full_name')
                        and not heir.get('heir_is_minor')):
                    declinations.append({
                        'name': heir['heir_full_name'],
                        'relationship': heir['heir_relationship'],
                        'gender': heir.get('heir_gender', 'Male'),
                    })

    return declinations


# --- Closing Document Selection (spec § 2.4) ---

def select_closing_documents(data):
    """Return list of (template_filename, output_title) for closing docs."""
    estate_type = data['estate_type']
    docs = []

    if estate_type == 'Testate':
        if data.get('all_heirs_sui_juris', True):
            docs.append(('Petition to Close Estate CURLY.docx',
                          'Petition to Close Estate'))
        else:
            docs.append(('Petition to Close Estate-No sui juris.docx',
                          'Petition to Close Estate'))
        docs.append(('Order to Close Estate CURLY.docx',
                      'Order to Close Estate'))
    else:
        docs.append(('Petition to Close Intestate Estate CURLY.docx',
                      'Petition to Close Intestate Estate'))
        docs.append(('Order Closing Intestate Estate CURLY.docx',
                      'Order Closing Intestate Estate'))

    return docs


# --- Receipt & Waiver Selection (spec § 2.5) ---

def select_receipt_waiver_template(heir, data):
    """Return (template_filename, output_title) for one heir's receipt & waiver."""
    estate_type = data['estate_type']
    pr_name = data.get('pr_full_name', '')

    if estate_type == 'Intestate':
        return ('Receipt and Waiver for Intestate Estate.docx',
                f'Receipt and Waiver - {heir["heir_full_name"]}')

    # Testate paths
    is_pr = heir['heir_full_name'] == pr_name
    btype = heir.get('heir_beneficiary_type', 'General')

    if is_pr and btype == 'Residuary':
        return ('Receipt and Waiver  - Residuary and Executor.docx',
                f'Receipt and Waiver - {heir["heir_full_name"]}')
    elif btype == 'Residuary':
        return ('Receipt and Waiver  - Residuary CURLY.docx',
                f'Receipt and Waiver - {heir["heir_full_name"]}')
    elif btype == 'Specific':
        return ('Receipt & Waiver - Testate CURLY.docx',
                f'Receipt and Waiver - {heir["heir_full_name"]}')
    else:
        return ('Receipt and Waiver  - General CURLY.docx',
                f'Receipt and Waiver - {heir["heir_full_name"]}')


# --- Flags & Warnings (spec § 2.6) ---

def generate_flags(data):
    """Return list of {level: 'warning'|'info', message: str}."""
    flags = []
    heirs = data.get('heirs', [])

    if data.get('pr_state', 'Tennessee') != 'Tennessee':
        flags.append({'level': 'warning',
                       'message': 'Nonresident PR — may not be eligible under § 30-1-116'})

    if data.get('pr_criminal_history') or data.get('pr_penitentiary_sentence'):
        flags.append({'level': 'warning',
                       'message': 'PR criminal history disclosed — attorney review required per § 30-1-117(10)'})

    est_val = data.get('estimated_estate_value', 0)
    if est_val and float(est_val) <= 50000 and not data.get('small_estate_election'):
        flags.append({'level': 'info',
                       'message': 'Estate may qualify for small estate administration under § 30-4-103'})

    property_counties = [p.get('property_county', '') for p in data.get('properties', [])]
    dec_county = data.get('decedent_county', '')
    out_of_county = [c for c in property_counties if c and c != dec_county]
    if out_of_county:
        flags.append({'level': 'warning',
                       'message': f'Real property in {", ".join(out_of_county)} — ancillary probate may be needed'})

    has_minor = any(h.get('heir_is_minor') for h in heirs)
    has_disability = any(h.get('heir_has_disability') for h in heirs)
    if has_minor or has_disability:
        flags.append({'level': 'warning',
                       'message': 'Not all heirs are sui juris — affects closing documents and may require guardian appointment'})

    if data.get('decedent_had_business'):
        flags.append({'level': 'info',
                       'message': 'Decedent had business interest — include details in petition per § 30-1-117(11)'})

    if (data.get('estate_type') == 'Testate'
            and not data.get('will_waives_inventory')):
        flags.append({'level': 'info',
                       'message': 'Will does not waive inventory — inventory will be required'})

    return flags


# --- Deadline Calculations (spec § 3.2) ---

def calculate_deadlines(dod_str, publication_date_str=None):
    """Calculate statutory deadlines.

    Args:
        dod_str: Date of death as YYYY-MM-DD
        publication_date_str: Publication date as YYYY-MM-DD (optional)

    Returns:
        dict of deadline names -> YYYY-MM-DD strings
    """
    dod = datetime.strptime(dod_str, '%Y-%m-%d')
    deadlines = {}

    # Absolute bar date: DOD + 12 months
    try:
        bar_date = dod.replace(year=dod.year + 1)
    except ValueError:
        # Handle leap year edge case (Feb 29 -> Feb 28)
        bar_date = dod.replace(year=dod.year + 1, day=28)
    deadlines['absolute_bar_date'] = bar_date.strftime('%Y-%m-%d')

    if publication_date_str:
        pub = datetime.strptime(publication_date_str, '%Y-%m-%d')

        # Claims deadline: publication + 4 months
        month = pub.month + 4
        year = pub.year
        if month > 12:
            month -= 12
            year += 1
        try:
            claims = pub.replace(year=year, month=month)
        except ValueError:
            import calendar
            last_day = calendar.monthrange(year, month)[1]
            claims = pub.replace(year=year, month=month, day=min(pub.day, last_day))
        deadlines['claims_deadline'] = claims.strftime('%Y-%m-%d')

        # Exception deadline: claims + 30 days
        exception = claims + timedelta(days=30)
        deadlines['exception_deadline'] = exception.strftime('%Y-%m-%d')

    return deadlines


# --- Build Common Replacements ---

def build_common_replacements(data):
    """Build the merge field replacement dict from intake data.

    Maps standardized spec field names AND legacy template field names
    to data values. This handles the inconsistent naming across templates.
    """
    dec_pronouns = derive_pronouns(data.get('decedent_gender', 'Male'))
    pr_pronouns = derive_pronouns(data.get('pr_gender', 'Male'))
    pr_title = derive_pr_title(data.get('estate_type', 'Testate'),
                                data.get('pr_gender', 'Male'))

    today = datetime.now()
    today_str = today.strftime('%Y-%m-%d')

    # Criminal history statement
    if data.get('pr_criminal_history') or data.get('pr_penitentiary_sentence'):
        criminal_stmt = data.get('pr_criminal_details',
                                  'has disclosed criminal history — see attached')
    else:
        criminal_stmt = ('has never been convicted of any felony or misdemeanor '
                         'and has never served any sentence of imprisonment in a penitentiary')

    # Business interest statement
    if data.get('decedent_had_business'):
        business_stmt = (f'was {data.get("decedent_business_details", "")} '
                         'the owner of or had a controlling interest in an ongoing '
                         'business or economic enterprise that is or may be part '
                         'of the estate to be administered')
    else:
        business_stmt = ('was not the owner of nor had a controlling interest in '
                         'any ongoing business or economic enterprise that is or '
                         'may be part of the estate to be administered')

    # Build heirs list block
    heirs = data.get('heirs', [])
    heirs_lines = []
    for h in heirs:
        age_str = f", Age {h.get('heir_age', 'unknown')}" if h.get('heir_age') else ''
        heirs_lines.append(
            f"{h['heir_full_name']}{age_str} — {h.get('heir_relationship', '')}\n"
            f"{h.get('heir_address', '')}, {h.get('heir_city', '')}"
        )
    heirs_list = '\n\n'.join(heirs_lines)

    # All heirs sui juris?
    all_sui_juris = all(
        not h.get('heir_is_minor') and not h.get('heir_has_disability')
        for h in heirs
    ) if heirs else True

    # Will waiver statement
    waivers = []
    if data.get('will_waives_bond'):
        waivers.append('making bond')
    if data.get('will_waives_inventory'):
        waivers.append('filing an inventory')
    if data.get('will_waives_accountings'):
        waivers.append('filing accountings')
    will_waiver_stmt = 'excused from ' + ', '.join(waivers) if waivers else ''

    dod_formatted = format_date_legal(data['decedent_dod']) if data.get('decedent_dod') else ''

    # Sui juris statement
    if all_sui_juris:
        sui_juris_stmt = 'All beneficiaries are sui juris.'
    else:
        sui_juris_stmt = 'Not all beneficiaries are sui juris.'

    # Computed closing-doc field: "were/were no" objections
    were_or_no = data.get('were_or_no_objections', 'were no')

    # Attorney first name
    attorney_full = data.get('attorney_full_name', '')
    attorney_first = attorney_full.split()[0] if attorney_full else ''

    # PR city + state + zip combined
    pr_city_state_zip = ', '.join(filter(None, [
        data.get('pr_city', ''),
        data.get('pr_state', 'Tennessee'),
        data.get('pr_zip', ''),
    ]))

    # Will execution date formatted
    will_date_formatted = (format_date_legal(data['will_execution_date'])
                           if data.get('will_execution_date') else '')

    # Codicil date
    codicil_date_formatted = (format_date_legal(data['codicil_execution_date'])
                              if data.get('codicil_execution_date') else '')

    replacements = {
        # --- Decedent (all template variants) ---
        '{DECEDENT NAME}': data.get('decedent_full_name', ''),
        '{DECEDENT}': data.get('decedent_full_name', ''),
        '{Decedent name}': data.get('decedent_full_name', ''),
        '{Decedent Name}': data.get('decedent_full_name', ''),
        '{Decedent}': data.get('decedent_full_name', ''),
        '{decedent name}': data.get('decedent_full_name', ''),
        '{DECEDENT ADDRESS}': data.get('decedent_address', ''),
        '{ADDRESS}': data.get('decedent_address', ''),
        '{address}': data.get('decedent_address', ''),
        '{Address}': data.get('decedent_address', ''),
        "{Decedent's Address}": data.get('decedent_address', ''),
        "{Decedent's Street Address}": data.get('decedent_address', ''),
        '{CITY}': data.get('decedent_city', ''),
        '{City}': data.get('decedent_city', ''),
        '{Decedent City}': data.get('decedent_city', ''),
        '{DECEDENT CITY}': data.get('decedent_city', ''),
        '{COUNTY}': data.get('decedent_county', ''),
        '{COUNTY NAME}': data.get('decedent_county', ''),
        '{County}': data.get('decedent_county', ''),
        '{DECEDENT COUNTY}': data.get('decedent_county', ''),
        '{County of Residence}': data.get('decedent_county', ''),
        '{County of Residence for Decedent}': data.get('decedent_county', ''),
        "{Decedent's County of Residence}": data.get('decedent_county', ''),
        '{County of Probate}': data.get('decedent_county', ''),
        '{DATE OF DEATH}': dod_formatted,
        '{Date of Death}': dod_formatted,
        "{Decedent's Date of Death}": dod_formatted,
        '{AGE OF DECEDENT}': str(data.get('decedent_age', '')),
        '{AGE}': str(data.get('decedent_age', '')),
        '{AGE AT DEATH}': str(data.get('decedent_age', '')),
        '{Age at Death}': str(data.get('decedent_age', '')),
        '{Age of Death}': str(data.get('decedent_age', '')),
        "{Decedent's Age}": str(data.get('decedent_age', '')),
        '{PLACE OF DEATH}': data.get('decedent_place_of_death', ''),
        '{Place of Death}': data.get('decedent_place_of_death', ''),
        '{At}': data.get('decedent_place_of_death', ''),
        '{HIS/HER}': dec_pronouns['possessive'],
        '{his/her}': dec_pronouns['possessive'],
        '{DECEDENT PRONOUN \u2013 HIS/HER}': dec_pronouns['possessive'],
        '{Decedent Possessive Pronoun}': dec_pronouns['possessive'],
        '{HE/SHE}': dec_pronouns['subject'],
        '{he/she}': dec_pronouns['subject'],
        '{him/her}': dec_pronouns['object'],
        '{was/was not CHOOSE ONE}': 'was' if data.get('decedent_had_business') else 'was not',
        '{Decedent Spouse}': data.get('decedent_spouse_name', ''),
        "{Decedent's Spouse Date of Death}": data.get('decedent_spouse_dod', ''),

        # --- Personal Representative / Petitioner (all template variants) ---
        '{PETITIONER NAME}': data.get('pr_full_name', ''),
        '{PETITIONER}': data.get('pr_full_name', ''),
        '{Petitioner name}': data.get('pr_full_name', ''),
        '{Petitioner Name}': data.get('pr_full_name', ''),
        '{Petitioner}': data.get('pr_full_name', ''),
        '{petitioner}': data.get('pr_full_name', ''),
        "{petitioner's name}": data.get('pr_full_name', ''),
        '{EXECUTOR NAME}': data.get('pr_full_name', ''),
        '{Executor Name}': data.get('pr_full_name', ''),
        '{Executor}': data.get('pr_full_name', ''),
        '{Administrator Name}': data.get('pr_full_name', ''),
        '{Administrator}': data.get('pr_full_name', ''),
        '{PETITIONER ADDRESS}': data.get('pr_address', ''),
        '{ADDRESS OF PETITIONER}': data.get('pr_address', ''),
        '{STREET OF PETITIONER}': data.get('pr_address', ''),
        '{Street Address of Petitioner}': data.get('pr_address', ''),
        '{Petitioner Street Address}': data.get('pr_address', ''),
        '{Petitioner Address}': data.get('pr_address', ''),
        '{PETITIONER CITY}': data.get('pr_city', ''),
        '{CITY OF PETITIONER}': data.get('pr_city', ''),
        '{City of Petitioner}': data.get('pr_city', ''),
        '{Petitioner City}': data.get('pr_city', ''),
        '{Petitioner City, State Zip}': pr_city_state_zip,
        '{Petitioner City, State, Zip}': pr_city_state_zip,
        '{Petitioner Zip}': data.get('pr_zip', ''),
        '{Petitioner Phone #}': data.get('pr_phone', ''),
        '{PETITIONER AGE}': str(data.get('pr_age', '')),
        '{AGE OF PETITIONER}': str(data.get('pr_age', '')),
        '{Age of Petitioner}': str(data.get('pr_age', '')),
        '{Petitioner Age}': str(data.get('pr_age', '')),
        '{RELATIONSHIP OF PETITIONER TO DECEDENT}': data.get('pr_relationship', ''),
        '{Relationship of Petitioner to Decedent}': data.get('pr_relationship', ''),
        '{Petitioner Relationship to Decedent}': data.get('pr_relationship', ''),
        '{Relation of Petition to Decedent}': data.get('pr_relationship', ''),
        "{Decedent's Relationship to Petitioner}": data.get('pr_relationship', ''),
        "{PETITIONER'S RELATIONSHIP TO THE DE}": data.get('pr_relationship', ''),
        '{RELATION TO DECEDENT}': data.get('pr_relationship', ''),

        # --- PR Title variants ---
        '{Executor/Executrix/PR Title CHOOSE ONE}': pr_title,
        '{Executor/Executrix/PR Title}': pr_title,
        '{Executor/Executrix/Personal Representative}': pr_title,
        '{Executor/Executrix}': pr_title,
        '{Executor/trix}': pr_title,
        '{Administrator/trix CHOOSE ONE}': pr_title,
        '{Administrator/trix}': pr_title,
        '{Administrator or Administratrix CHOOSE ONE}': pr_title,
        '{Administrator or Administratrix}': pr_title,
        '{administrator or administratrix}': pr_title.lower(),
        '{Administrator/Executor/PR}': pr_title,
        '{Title Executor/Executrix CHOOSE ONE}': pr_title,
        '{TITLE}': pr_title,
        '{Title}': pr_title,
        '{PETITIONER PRONOUNT \u2013 HE/SHE}': pr_pronouns['subject'],
        '{PETITIONER PRONOUN \u2013 HE/SHE}': pr_pronouns['subject'],

        # --- Will details ---
        '{Date of LWT}': format_date_legal(data['will_execution_date']) if data.get('will_execution_date') else '',
        '{WITNESS 1}': data.get('will_witness_1', ''),
        '{WITNESS 2}': data.get('will_witness_2', ''),
        '{PARAGRAPH # APPOINTING PETITIONER}': data.get('will_appointment_paragraph', ''),
        '{Paragraph/Article/etc}': data.get('will_appointment_paragraph', ''),

        # --- Heirs ---
        '{HEIRS_LIST}': heirs_list,

        # --- Case / Court ---
        '{STATE}': data.get('decedent_state', 'Tennessee'),
        '{Docket Number}': data.get('case_number', ''),
        '{MONTH}': today.strftime('%B'),
        '{CURRENT YEAR}': str(today.year),
        '{current_date_day}': ordinal_day(today_str),

        # --- Computed statements ---
        '{CRIMINAL_STATEMENT}': criminal_stmt,
        '{BUSINESS_STATEMENT}': business_stmt,
        '{WAIVER_STATEMENT}': will_waiver_stmt,
        '{SUI_JURIS_STATEMENT}': sui_juris_stmt,
        '{were/were no}': were_or_no,

        # --- Additional case/date variants ---
        '{docket number}': data.get('case_number', ''),
        '{YEAR}': str(today.year),
        '{Current Year}': str(today.year),
        '{Month}': today.strftime('%B'),

        # --- Additional attorney variants ---
        '{ATTORNEY}': data.get('attorney_full_name', ''),
        '{Attorney first name}': attorney_first,
        '{BPR #}': data.get('attorney_bpr', ''),

        # --- Additional PR / petitioner pronoun variants ---
        '{Petitioner Pronoun HIS/HER}': pr_pronouns['possessive'],
        '{Petitioner Pronoun HE/SHE}': pr_pronouns['subject'],

        # --- Additional relationship variants ---
        '{RELATION OF INHERITORS \u2013 sister, brother, children, etc}': data.get('pr_relationship', ''),

        # --- Will date variants ---
        '{Will Execution Date}': will_date_formatted,
        '{WILL DATE}': will_date_formatted,
        '{Codicil Date}': codicil_date_formatted,

        # --- Firm ---
        '{ATTORNEY NAME}': data.get('attorney_full_name', ''),
        '{BPR}': data.get('attorney_bpr', ''),
        'Dale, Hutto & Lyle, PLLC': data.get('firm_name', 'Muletown Law, P.C.'),
    }

    return replacements
