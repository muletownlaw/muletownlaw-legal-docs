from http.server import BaseHTTPRequestHandler
import json
from io import BytesIO
from docx import Document
import urllib.request
import sys
import os
from datetime import datetime
import re

# Add the api directory to path so we can import template_config
sys.path.insert(0, os.path.dirname(__file__))

def format_name_for_filename(full_name):
    """Format name as 'Lastname Firstname' removing middle initials"""
    name_parts = full_name.strip().split()
    if len(name_parts) < 2:
        return full_name

    # Remove middle initials (single letter followed by optional period)
    filtered_parts = [part for part in name_parts if not re.match(r'^[A-Z]\.?$', part)]

    if len(filtered_parts) >= 2:
        # Last name is last element, first name is everything before it
        lastname = filtered_parts[-1]
        firstname = ' '.join(filtered_parts[:-1])
        return f"{lastname} {firstname}"
    else:
        return full_name

try:
    from template_config import TEMPLATE_URLS
    print("[ACP] Successfully imported TEMPLATE_URLS")
    print(f"[ACP] ACP template URL: {TEMPLATE_URLS.get('acp', 'NOT FOUND')}")
except ImportError as e:
    print(f"[ACP] CRITICAL: Failed to import template_config: {e}")
    TEMPLATE_URLS = {'acp': 'ERROR_NO_CONFIG'}

# Module-level cache to avoid re-downloading templates on warm invocations
_template_cache = {}

def download_template(url):
    """Download template from Google Drive, with in-memory caching.

    Uses requests for reliable redirect and cookie handling. Google Drive
    sometimes returns an HTML confirmation page for larger files; this
    function detects that and retries with the embedded confirm token.
    """
    if url in _template_cache:
        print(f"[ACP] Using cached template for: {url}")
        return BytesIO(_template_cache[url])
    try:
        import requests as req_lib
        print(f"[ACP] Downloading template from: {url}")
        session = req_lib.Session()
        session.headers['User-Agent'] = 'Mozilla/5.0'
        response = session.get(url, allow_redirects=True, timeout=30)
        response.raise_for_status()
        content = response.content
        if content[:4] != b'PK\x03\x04':
            confirm = re.search(rb'confirm=([0-9A-Za-z_\-]+)', content)
            file_id = re.search(r'[?&]id=([^&]+)', url)
            if confirm and file_id:
                retry_url = (
                    f'https://drive.google.com/uc?export=download'
                    f'&id={file_id.group(1)}'
                    f'&confirm={confirm.group(1).decode()}'
                )
                response = session.get(retry_url, allow_redirects=True, timeout=30)
                content = response.content
        if content[:4] != b'PK\x03\x04':
            raise Exception("Downloaded file is not a valid .docx (failed ZIP header check). "
                            "Ensure the file is shared as 'Anyone with the link can view'.")
        print(f"[ACP] Template downloaded: {len(content)} bytes")
        _template_cache[url] = content
        return BytesIO(content)
    except Exception as e:
        print(f"[ACP] Template download failed: {e}")
        raise Exception(f"Failed to download template: {str(e)}")

def merge_runs_in_paragraph(paragraph):
    """Merge all runs in a paragraph to handle split placeholders"""
    if not paragraph.runs:
        return
    
    # Get full text
    full_text = paragraph.text
    
    # Clear all runs
    for run in paragraph.runs:
        run.text = ''
    
    # Put all text in first run, preserving its formatting
    if paragraph.runs:
        paragraph.runs[0].text = full_text

def replace_in_document(doc, replacements):
    """Replace all placeholders - handles split runs"""
    
    # First, merge runs in all paragraphs to consolidate text
    for paragraph in doc.paragraphs:
        merge_runs_in_paragraph(paragraph)
    
    # Also merge runs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    merge_runs_in_paragraph(paragraph)
    
    # Now do replacements (runs are merged, so text is in single runs)
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)

def generate_acp_document(data):
    """Generate ACP from Google Drive template"""

    # Get template URL from config
    template_url = TEMPLATE_URLS.get('acp', '')

    if not template_url or template_url == 'ERROR_NO_CONFIG':
        raise Exception("Template configuration not found. Check that template_config.py is in api/ folder")

    # Remove trailing slash if present
    template_url = template_url.rstrip('/')

    print(f"[ACP] Using template URL: {template_url}")

    # Download template from Google Drive
    template_buffer = download_template(template_url)

    # Open the template
    doc = Document(template_buffer)
    
    pronoun = data.get('CLIENT_PRONOUN', 'he' if data.get('CLIENT_GENDER') == 'Male' else 'she')
    
    replacements = {
        '{CLIENT_NAME}': data['CLIENT_NAME'].upper(),
        '{CLIENT_PRONOUN}': pronoun,
        '{PRIMARY_AGENT_NAME}': data['PRIMARY_AGENT_NAME'].upper(),
        '{PRIMARY_AGENT_RELATION}': data['PRIMARY_AGENT_RELATION'],
        '{ALTERNATE_AGENT_NAME}': data['ALTERNATE_AGENT_NAME'].upper(),
        '{ALTERNATE_AGENT_RELATION}': data['ALTERNATE_AGENT_RELATION'],
        '{EXEC_MONTH}': data.get('EXEC_MONTH', 'October'),
        '{EXEC_YEAR}': data.get('EXEC_YEAR', '2025'),
    }
    
    replace_in_document(doc, replacements)
    return doc

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            doc = generate_acp_document(data)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Format filename as: YYYY-MM-DD ACP lastname firstname.docx
            today = datetime.now().strftime('%Y-%m-%d')
            formatted_name = format_name_for_filename(data["CLIENT_NAME"])
            filename = f"{today} ACP {formatted_name}.docx"

            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.end_headers()
            self.wfile.write(buffer.getvalue())
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
