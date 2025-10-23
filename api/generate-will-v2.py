from http.server import BaseHTTPRequestHandler
import json
from io import BytesIO
from docx import Document
import os
import re

def merge_runs_in_paragraph(paragraph):
    """Merge all runs to handle split placeholders"""
    if not paragraph.runs:
        return
    full_text = paragraph.text
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = full_text

def replace_in_document(doc, replacements):
    """Replace all placeholders in document"""
    for paragraph in doc.paragraphs:
        merge_runs_in_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    merge_runs_in_paragraph(paragraph)
    
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)

def load_clause_content(clause_filename):
    """Load clause content as text, skipping markers"""
    clause_path = os.path.join(os.path.dirname(__file__), 'clauses', clause_filename)
    try:
        clause_doc = Document(clause_path)
        lines = []
        for para in clause_doc.paragraphs:
            text = para.text.strip()
            if text and not text.startswith('##'):
                lines.append(text)
        return '\n\n'.join(lines)
    except:
        return ''

def format_children_list(children):
    """Format children array into text"""
    if not children:
        return ""
    
    if len(children) == 1:
        child = children[0]
        return f"{child['name']}, born {child['dob']}"
    
    result = []
    for i, child in enumerate(children):
        if i == len(children) - 1:
            result.append(f"and {child['name']}, born {child['dob']}")
        else:
            result.append(f"{child['name']}, born {child['dob']}")
    
    return ", ".join(result)

def number_to_words(num):
    """Convert number to words"""
    words = ["one", "two", "three", "four", "five", "six", "seven", "eight", "nine", "ten"]
    if num <= 10:
        return words[num - 1]
    return str(num)

def generate_will_document(data):
    """Generate will using marker-based template"""
    template_path = os.path.join(os.path.dirname(__file__), 'Simple_Will.docx')
    doc = Document(template_path)
    
    # Calculate derived values
    client_gender = data.get('CLIENT_GENDER', 'Male')
    is_married = bool(data.get('CLIENT_SPOUSE_NAME', '').strip())
    spouse_gender = 'Female' if client_gender == 'Male' else 'Male'
    
    # Pronouns
    if client_gender == 'Male':
        client_pronoun_subj = 'he'
        client_pronoun_poss = 'his'
        testator_title = 'Testator'
        executor_title = 'Executor'
        spouse_type = 'wife'
    else:
        client_pronoun_subj = 'she'
        client_pronoun_poss = 'her'
        testator_title = 'Testatrix'
        executor_title = 'Executrix'
        spouse_type = 'husband'
    
    if spouse_gender == 'Male':
        spouse_pronoun_subj = 'he'
        spouse_pronoun_poss = 'his'
    else:
        spouse_pronoun_subj = 'she'
        spouse_pronoun_poss = 'her'
    
    # Children
    children = data.get('children', [])
    num_children = len(children)
    children_list = format_children_list(children)
    child_or_children = 'child' if num_children == 1 else 'children'
    
    # Build replacements
    replacements = {
        '{CLIENT_NAME}': data.get('CLIENT_NAME', '').upper(),
        '{CLIENT_COUNTY}': data.get('CLIENT_COUNTY', 'Maury'),
        '{CLIENT_PRONOUN_SUBJECTIVE}': client_pronoun_subj,
        '{CLIENT_PRONOUN_POSSESSIVE}': client_pronoun_poss,
        '{CLIENT_SPOUSE_NAME}': data.get('CLIENT_SPOUSE_NAME', '').upper() if is_married else '',
        '{SPOUSE_TYPE}': spouse_type,
        '{SPOUSE_PRONOUN_SUBJECTIVE}': spouse_pronoun_subj,
        '{SPOUSE_PRONOUN_POSSESSIVE}': spouse_pronoun_poss,
        '{TESTATOR_TITLE}': testator_title,
        '{EXECUTOR_TITLE}': executor_title,
        '{ALTERNATE_EXECUTOR_NAME}': data.get('ALTERNATE_EXECUTOR_NAME', '').upper(),
        '{ALTERNATE_EXECUTOR_RELATION}': data.get('ALTERNATE_EXECUTOR_RELATION', ''),
        '{ALTERNATE_EXECUTOR_COUNTY}': data.get('ALTERNATE_EXECUTOR_COUNTY', data.get('CLIENT_COUNTY', 'Maury')),
        '{ALTERNATE_EXECUTOR_STATE}': data.get('ALTERNATE_EXECUTOR_STATE', 'Tennessee'),
        '{EXEC_MONTH}': data.get('EXEC_MONTH', 'October'),
        '{EXEC_YEAR}': data.get('EXEC_YEAR', '2025'),
        '{NUMBER_OF_CHILDREN}': number_to_words(num_children) if num_children > 0 else 'no',
        '{CHILDREN_LIST}': children_list,
        '{CHILD_OR_CHILDREN}': child_or_children,
        '{CONTINGENT_BENEFICIARY_NAME}': data.get('CONTINGENT_BENEFICIARY_NAME', '').upper(),
        '{CONTINGENT_BENEFICIARY_RELATION}': data.get('CONTINGENT_BENEFICIARY_RELATION', ''),
        '{DISINHERITED_NAME}': data.get('DISINHERITED_NAME', '').upper() if data.get('INCLUDE_DISINHERITANCE') else '',
        '{DISINHERITED_RELATION}': data.get('DISINHERITED_RELATION', '') if data.get('INCLUDE_DISINHERITANCE') else '',
    }
    
    # Step 1: Replace all variables FIRST
    replace_in_document(doc, replacements)
    
    # Step 2: Handle conditional sections by modifying text, not removing paragraphs
    for para in doc.paragraphs:
        text = para.text
        
        # Handle ##IF_MARRIED## sections
        if '##IF_MARRIED##' in text:
            if is_married:
                # Remove markers, keep content
                new_text = text.replace('##IF_MARRIED##', '').replace('##END_IF##', '')
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = new_text
            else:
                # Replace entire paragraph with empty
                for run in para.runs:
                    run.text = ''
        
        # Handle ##Delete first sentence if unmarried##
        elif '##Delete first sentence if unmarried##' in text:
            if is_married:
                # Remove marker only
                new_text = text.replace('##Delete first sentence if unmarried##', '')
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = new_text
            else:
                # Keep only the children part
                if 'I have' in text:
                    start_idx = text.find('I have')
                    new_text = text[start_idx:].replace('##Delete first sentence if unmarried##', '')
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = new_text
    
    # Step 3: Build clause insertions
    article_iii_clauses = []
    if data.get('INCLUDE_HANDWRITTEN_LIST'):
        content = load_clause_content('Handwritten_List.docx')
        if content:
            article_iii_clauses.append(content)
    
    if data.get('INCLUDE_DISINHERITANCE'):
        content = load_clause_content('Love_And_Affection.docx')
        if content:
            article_iii_clauses.append(content)
    
    if data.get('INCLUDE_REAL_ESTATE_DEBT'):
        content = load_clause_content('Real_Estate_Debt.docx')
        if content:
            article_iii_clauses.append(content)
    
    # New articles
    new_articles = []
    article_counter = 6
    
    if data.get('INCLUDE_NO_CONTEST'):
        content = load_clause_content('No_Contest.docx')
        if content:
            new_articles.append(f"Article {article_counter} - No Contest\n\n{content}")
            article_counter += 1
    
    # Step 4: Replace insertion markers with clause content
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if text == '##INSERT_ARTICLE_III_CLAUSES##':
            if article_iii_clauses:
                replacement = '\n\n'.join(article_iii_clauses)
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = replacement
            else:
                # Empty the marker
                for run in para.runs:
                    run.text = ''
        
        elif text == '##INSERT_NEW_ARTICLES##':
            if new_articles:
                replacement = '\n\n'.join(new_articles)
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = replacement
            else:
                # Empty the marker
                for run in para.runs:
                    run.text = ''
    
    # Step 5: Renumber articles if needed
    if new_articles:
        # Find where the new articles were inserted and renumber from there
        found_new_article = False
        current_article_num = article_counter  # Continue from where we left off
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Check if this paragraph contains our inserted articles
            if 'Article 6 - No Contest' in text:
                found_new_article = True
                continue
            
            # After finding new articles, renumber subsequent articles
            if found_new_article and text.startswith('Article '):
                match = re.match(r'Article (\w+) - (.+)', text)
                if match:
                    title = match.group(2)
                    for run in para.runs:
                        if 'Article' in run.text:
                            run.text = f'Article {current_article_num} - {title}'
                            current_article_num += 1
                            break
    
    # Step 6: Fix unmarried executor appointments and spouse references
    if not is_married:
        for para in doc.paragraphs:
            text = para.text
            if text.strip():  # Only process non-empty paragraphs
                # Fix executor appointment
                if 'appoint my' in text and ', ,' in text:
                    new_text = text.replace(
                        f'my {spouse_type}, ,',
                        replacements['{ALTERNATE_EXECUTOR_NAME}'] + ','
                    )
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = new_text
                
                # Fix "my said wife/husband" references
                elif f'my said {spouse_type}, ,' in text:
                    new_text = text.replace(f'my said {spouse_type}, ,', 'my children')
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = new_text
                
                # Fix general spouse references with blank
                elif f'{spouse_type}, ,' in text:
                    # Remove this paragraph's content if it's just about spouse
                    if 'survives me' in text and 'children' not in text:
                        for run in para.runs:
                            run.text = ''
    
    # Step 7: Final cleanup - remove any remaining markers
    for para in doc.paragraphs:
        text = para.text
        if '##' in text or (not is_married and ', ,' in text):
            for run in para.runs:
                # Remove any remaining ## markers
                cleaned = run.text.replace('##', '')
                # Fix any remaining blank spouse references
                if not is_married:
                    cleaned = cleaned.replace(f'my {spouse_type}, ,', replacements.get('{ALTERNATE_EXECUTOR_NAME}', 'alternate executor'))
                    cleaned = cleaned.replace(f'my said {spouse_type}, ,', 'my children')
                    cleaned = cleaned.replace(f'{spouse_type}, ,', '')
                run.text = cleaned
    
    return doc

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            doc = generate_will_document(data)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            client_name = data.get('CLIENT_NAME', 'Client').replace(' ', '_')
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="Will_{client_name}.docx"')
            self.end_headers()
            self.wfile.write(buffer.getvalue())
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            error_msg = f"{type(e).__name__}: {str(e)}"
            self.wfile.write(json.dumps({'error': error_msg}).encode())
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
