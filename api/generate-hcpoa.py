from http.server import BaseHTTPRequestHandler
import json
from io import BytesIO
from docx import Document
import urllib.request
import sys
import os
from datetime import datetime
import re

# Add the api directory to path so we can import template_config
sys.path.insert(0, os.path.dirname(__file__))

def format_name_for_filename(full_name):
    """Format name as 'Lastname Firstname' removing middle initials"""
    name_parts = full_name.strip().split()
    if len(name_parts) < 2:
        return full_name

    # Remove middle initials (single letter followed by optional period)
    filtered_parts = [part for part in name_parts if not re.match(r'^[A-Z]\.?$', part)]

    if len(filtered_parts) >= 2:
        # Last name is last element, first name is everything before it
        lastname = filtered_parts[-1]
        firstname = ' '.join(filtered_parts[:-1])
        return f"{lastname} {firstname}"
    else:
        return full_name

try:
    from template_config import TEMPLATE_URLS
    print("[HCPOA] Successfully imported TEMPLATE_URLS")
    print(f"[HCPOA] HCPOA template URL: {TEMPLATE_URLS.get('hcpoa', 'NOT FOUND')}")
except ImportError as e:
    print(f"[HCPOA] CRITICAL: Failed to import template_config: {e}")
    TEMPLATE_URLS = {'hcpoa': 'ERROR_NO_CONFIG'}

def download_template(url):
    """Download template from Google Drive"""
    try:
        print(f"[HCPOA] Downloading template from: {url}")
        with urllib.request.urlopen(url) as response:
            template_bytes = response.read()
            print(f"[HCPOA] Template downloaded: {len(template_bytes)} bytes")
            return BytesIO(template_bytes)
    except Exception as e:
        print(f"[HCPOA] Template download failed: {e}")
        raise Exception(f"Failed to download template: {str(e)}")

def replace_in_document(doc, replacements):
    """Replace all placeholders in the document"""
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)

def generate_hcpoa_document(data):
    """Generate HCPOA from Google Drive template"""

    # Get template URL from config
    template_url = TEMPLATE_URLS.get('hcpoa', '')

    if not template_url or template_url == 'ERROR_NO_CONFIG':
        raise Exception("Template configuration not found. Check that template_config.py is in api/ folder")

    # Remove trailing slash if present
    template_url = template_url.rstrip('/')

    print(f"[HCPOA] Using template URL: {template_url}")

    # Download template from Google Drive
    template_buffer = download_template(template_url)

    # Open the template
    doc = Document(template_buffer)
    
    replacements = {
        '{CLIENT_NAME}': data['CLIENT_NAME'].upper(),
        '{CLIENT_GENDER}': data.get('CLIENT_GENDER', 'Male'),
        '{CLIENT_COUNTY}': data['CLIENT_COUNTY'],
        '{PRIMARY_AGENT_NAME}': data['PRIMARY_AGENT_NAME'].upper(),
        '{PRIMARY_AGENT_RELATION}': data['PRIMARY_AGENT_RELATION'],
        '{PRIMARY_AGENT_COUNTY}': data.get('PRIMARY_AGENT_COUNTY', data['CLIENT_COUNTY']),
        '{ALTERNATE_AGENT_NAME}': data['ALTERNATE_AGENT_NAME'].upper(),
        '{ALTERNATE_AGENT_RELATION}': data['ALTERNATE_AGENT_RELATION'],
        '{ALTERNATE_AGENT_COUNTY}': data.get('ALTERNATE_AGENT_COUNTY', data['CLIENT_COUNTY']),
        '{EXEC_MONTH}': data.get('EXEC_MONTH', 'October'),
        '{EXEC_YEAR}': data.get('EXEC_YEAR', '2025')
    }
    
    replace_in_document(doc, replacements)
    return doc

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            doc = generate_hcpoa_document(data)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Format filename as: YYYY-MM-DD HCPOA lastname firstname.docx
            today = datetime.now().strftime('%Y-%m-%d')
            formatted_name = format_name_for_filename(data["CLIENT_NAME"])
            filename = f"{today} HCPOA {formatted_name}.docx"

            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.end_headers()
            self.wfile.write(buffer.getvalue())
            
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
