from http.server import BaseHTTPRequestHandler
import json
from io import BytesIO
from docx import Document
import urllib.request
import sys
import os
import traceback
from datetime import datetime
import re

# Add the api directory to path so we can import template_config
sys.path.insert(0, os.path.dirname(__file__))

def format_name_for_filename(full_name):
    """Format name as 'Lastname Firstname' removing middle initials"""
    name_parts = full_name.strip().split()
    if len(name_parts) < 2:
        return full_name

    # Remove middle initials (single letter followed by optional period)
    filtered_parts = [part for part in name_parts if not re.match(r'^[A-Z]\.?$', part)]

    if len(filtered_parts) >= 2:
        # Last name is last element, first name is everything before it
        lastname = filtered_parts[-1]
        firstname = ' '.join(filtered_parts[:-1])
        return f"{lastname} {firstname}"
    else:
        return full_name

try:
    from template_config import TEMPLATE_URLS
    print("[POA] Successfully imported TEMPLATE_URLS")
    print(f"[POA] POA template URL: {TEMPLATE_URLS.get('poa', 'NOT FOUND')}")
except ImportError as e:
    print(f"[POA] CRITICAL: Failed to import template_config: {e}")
    print(f"[POA] Current directory: {os.getcwd()}")
    print(f"[POA] Files in current dir: {os.listdir('.')}")
    TEMPLATE_URLS = {'poa': 'ERROR_NO_CONFIG'}

# Module-level cache to avoid re-downloading templates on warm invocations
_template_cache = {}

def download_template(url):
    """Download template from Google Drive, with in-memory caching.

    Uses requests for reliable redirect and cookie handling. Google Drive
    sometimes returns an HTML confirmation page for larger files; this
    function detects that and retries with the embedded confirm token.
    """
    if url in _template_cache:
        print(f"[POA] Using cached template for: {url}")
        return BytesIO(_template_cache[url])
    try:
        import requests as req_lib
        print(f"[POA] Downloading template from: {url}")
        session = req_lib.Session()
        session.headers['User-Agent'] = 'Mozilla/5.0'
        response = session.get(url, allow_redirects=True, timeout=30)
        response.raise_for_status()
        content = response.content
        # If Google Drive returned a confirmation page instead of the file,
        # extract the confirm token and retry.
        if content[:4] != b'PK\x03\x04':
            confirm = re.search(rb'confirm=([0-9A-Za-z_\-]+)', content)
            file_id = re.search(r'[?&]id=([^&]+)', url)
            if confirm and file_id:
                retry_url = (
                    f'https://drive.google.com/uc?export=download'
                    f'&id={file_id.group(1)}'
                    f'&confirm={confirm.group(1).decode()}'
                )
                response = session.get(retry_url, allow_redirects=True, timeout=30)
                content = response.content
        if content[:4] != b'PK\x03\x04':
            raise Exception("Downloaded file is not a valid .docx (failed ZIP header check). "
                            "Ensure the file is shared as 'Anyone with the link can view'.")
        print(f"[POA] Template downloaded: {len(content)} bytes")
        _template_cache[url] = content
        return BytesIO(content)
    except Exception as e:
        print(f"[POA] Template download failed: {e}")
        raise Exception(f"Failed to download template: {str(e)}")

def replace_placeholders(doc, data):
    """Replace placeholders in the document with actual data"""

    # Create replacement map
    replacements = {
        '{CLIENT_NAME}': data['CLIENT_NAME'].upper(),
        '{CLIENT_COUNTY}': data.get('COUNTY', data.get('CLIENT_COUNTY', '')),
        '{PRIMARY_AGENT_NAME}': data.get('AIF_NAME', data.get('PRIMARY_AGENT_NAME', '')).upper(),
        '{PRIMARY_AGENT_RELATION}': data.get('AIF_RELATIONSHIP', data.get('PRIMARY_AGENT_RELATION', '')),
        '{PRIMARY_AGENT_COUNTY}': data.get('COUNTY', data.get('CLIENT_COUNTY', '')),
        '{ALTERNATE_AGENT_NAME}': data.get('ALTERNATE_AIF_NAME', data.get('ALTERNATE_AGENT_NAME', '')).upper(),
        '{ALTERNATE_AGENT_RELATION}': data.get('ALTERNATE_AIF_RELATIONSHIP', data.get('ALTERNATE_AGENT_RELATION', '')),
        '{ALTERNATE_AGENT_COUNTY}': data.get('COUNTY', data.get('CLIENT_COUNTY', '')),
        '{EXEC_MONTH}': data['EXEC_MONTH'].upper(),
        '{EXEC_YEAR}': data['EXEC_YEAR'],
        '{AttorneyName}': data.get('ATTORNEY_NAME', 'Thomas M. Hutto')
    }
    
    print(f"[POA] Replacing {len(replacements)} placeholders")
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                # Replace inline while preserving formatting
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
    
    # Replace in headers/footers
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
        # Footer
        for paragraph in section.footer.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
    
    print("[POA] Placeholder replacement complete")
    return doc

def generate_poa_document(data):
    """Generate POA document from Google Drive template"""
    
    # Get template URL from config
    template_url = TEMPLATE_URLS.get('poa', '')
    
    if not template_url or template_url == 'ERROR_NO_CONFIG':
        raise Exception("Template configuration not found. Check that template_config.py is in api/ folder")
    
    # Remove trailing slash if present
    template_url = template_url.rstrip('/')
    
    print(f"[POA] Using template URL: {template_url}")
    
    # Download template from Google Drive
    template_buffer = download_template(template_url)
    
    # Open the template
    doc = Document(template_buffer)
    
    # Replace all placeholders with actual data
    doc = replace_placeholders(doc, data)
    
    return doc

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            # Get request body
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data.decode('utf-8'))
            
            print(f"[POA] Received request with keys: {list(data.keys())}")
            
            # Generate document
            doc = generate_poa_document(data)
            
            # Save to BytesIO buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            print(f"[POA] Document generated: {len(buffer.getvalue())} bytes")

            # Format filename as: YYYY-MM-DD POA lastname firstname.docx
            today = datetime.now().strftime('%Y-%m-%d')
            formatted_name = format_name_for_filename(data["CLIENT_NAME"])
            filename = f"{today} POA {formatted_name}.docx"

            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.end_headers()
            self.wfile.write(buffer.getvalue())
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"[POA ERROR] {error_trace}")
            
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({
                'error': str(e),
                'type': type(e).__name__,
                'traceback': error_trace
            }).encode())
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
