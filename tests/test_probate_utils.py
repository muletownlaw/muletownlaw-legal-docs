# tests/test_probate_utils.py
import pytest
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'api'))
from probate_utils import (derive_pronouns, derive_pr_title,
                           select_opening_documents, determine_declinations,
                           calculate_deadlines, generate_flags, format_date_legal,
                           ordinal_day, select_closing_documents,
                           select_receipt_waiver_template, build_common_replacements,
                           merge_runs_in_paragraph, replace_in_document,
                           load_template, build_zip)


class TestDerivePronouns:
    def test_male_pronouns(self):
        result = derive_pronouns('Male')
        assert result == {'subject': 'he', 'possessive': 'his', 'object': 'him'}

    def test_female_pronouns(self):
        result = derive_pronouns('Female')
        assert result == {'subject': 'she', 'possessive': 'her', 'object': 'her'}


class TestDerivePRTitle:
    def test_testate_male(self):
        assert derive_pr_title('Testate', 'Male') == 'Executor'

    def test_testate_female(self):
        assert derive_pr_title('Testate', 'Female') == 'Executrix'

    def test_intestate_male(self):
        assert derive_pr_title('Intestate', 'Male') == 'Administrator'

    def test_intestate_female(self):
        assert derive_pr_title('Intestate', 'Female') == 'Administratrix'


class TestFormatDateLegal:
    def test_standard_date(self):
        assert format_date_legal('2026-04-10') == 'April 10, 2026'

    def test_single_digit_day(self):
        assert format_date_legal('2026-01-05') == 'January 5, 2026'


class TestOrdinalDay:
    def test_10th(self):
        assert ordinal_day('2026-04-10') == '10th'

    def test_1st(self):
        assert ordinal_day('2026-04-01') == '1st'

    def test_2nd(self):
        assert ordinal_day('2026-04-02') == '2nd'

    def test_3rd(self):
        assert ordinal_day('2026-04-03') == '3rd'

    def test_11th(self):
        assert ordinal_day('2026-04-11') == '11th'

    def test_12th(self):
        assert ordinal_day('2026-04-12') == '12th'

    def test_13th(self):
        assert ordinal_day('2026-04-13') == '13th'

    def test_21st(self):
        assert ordinal_day('2026-04-21') == '21st'


class TestMergeRunsInParagraph:
    def test_merges_split_runs(self):
        """merge_runs_in_paragraph should consolidate all run text into first run."""
        from docx import Document
        doc = Document()
        para = doc.add_paragraph('')
        # Simulate Word splitting a placeholder across runs
        run1 = para.add_run('{DECEDENT')
        run2 = para.add_run(' NAME}')
        assert len(para.runs) >= 2
        merge_runs_in_paragraph(para)
        # All text should be in first run
        assert para.runs[0].text == '{DECEDENT NAME}'
        for run in para.runs[1:]:
            assert run.text == ''

    def test_handles_empty_paragraph(self):
        """Should not crash on empty paragraphs."""
        from docx import Document
        doc = Document()
        para = doc.add_paragraph('')
        # Clear the default run
        for run in para.runs:
            run.text = ''
        merge_runs_in_paragraph(para)  # Should not raise


class TestReplaceInDocument:
    def test_replaces_in_paragraphs(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph('{DECEDENT NAME} lived in {COUNTY}.')
        replace_in_document(doc, {'{DECEDENT NAME}': 'John Smith', '{COUNTY}': 'Maury'})
        assert doc.paragraphs[0].text == 'John Smith lived in Maury.'

    def test_replaces_in_tables(self):
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].paragraphs[0].add_run('{DECEDENT NAME}')
        replace_in_document(doc, {'{DECEDENT NAME}': 'Jane Doe'})
        # After merge_runs + replace, the placeholder is fully replaced
        assert 'Jane Doe' in table.rows[0].cells[0].paragraphs[0].text


class TestLoadTemplate:
    def test_loads_existing_template(self):
        """Should load a real template from probate-templates/."""
        doc = load_template('Personal Representative Oath CURLY.docx')
        assert doc is not None
        # It should be a Document object
        from docx import Document
        assert hasattr(doc, 'paragraphs')

    def test_raises_on_missing_template(self):
        with pytest.raises(Exception):
            load_template('nonexistent_template.docx')


class TestBuildZip:
    def test_builds_zip_with_documents(self):
        import zipfile
        from io import BytesIO
        from docx import Document

        doc1 = Document()
        doc1.add_paragraph('Test document 1')
        doc2 = Document()
        doc2.add_paragraph('Test document 2')

        documents = [('Petition', doc1), ('Order', doc2)]
        result = build_zip(documents, '2026-04-10')

        assert isinstance(result, BytesIO)
        with zipfile.ZipFile(result, 'r') as zf:
            names = zf.namelist()
            assert '2026-04-10 Petition.docx' in names
            assert '2026-04-10 Order.docx' in names

    def test_defaults_to_today_date(self):
        import zipfile
        from datetime import datetime
        from docx import Document

        doc = Document()
        doc.add_paragraph('Test')
        result = build_zip([('Doc', doc)])

        today = datetime.now().strftime('%Y-%m-%d')
        with zipfile.ZipFile(result, 'r') as zf:
            names = zf.namelist()
            assert f'{today} Doc.docx' in names


class TestSelectOpeningDocuments:
    def test_testate_standard_witnessed(self):
        data = {'estate_type': 'Testate', 'will_type': 'Standard Witnessed'}
        docs = select_opening_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition to Probate Will Ltrs Testamentary CURLY (1).docx' in filenames
        assert 'Order to Probate LWT CURLY.docx' in filenames
        assert 'Personal Representative Oath CURLY.docx' in filenames

    def test_intestate(self):
        data = {'estate_type': 'Intestate'}
        docs = select_opening_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition for Appointment of Administrator CURLY.docx' in filenames
        assert 'Order for Intestate Administration.docx' in filenames

    def test_small_estate(self):
        data = {'estate_type': 'Testate', 'small_estate_election': True}
        docs = select_opening_documents(data)
        assert len(docs) == 2  # Affidavit + Order only, no oath

    def test_holographic(self):
        data = {'estate_type': 'Testate', 'will_type': 'Holographic'}
        docs = select_opening_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition to Probate Holographic Will.docx' in filenames

    def test_muniment(self):
        data = {'estate_type': 'Testate', 'muniment_only': True}
        docs = select_opening_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition for Muniment of Title.docx' in filenames

    def test_will_codicil(self):
        data = {'estate_type': 'Testate', 'will_type': 'Will + Codicil'}
        docs = select_opening_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition to Probate Will and Codicil.docx' in filenames
        assert 'Order Admitting Codicil and LWT.docx' in filenames
        assert 'Personal Representative Oath CURLY.docx' in filenames


class TestDetermineDeclinations:
    def test_testate_named_executor_can_serve(self):
        data = {'estate_type': 'Testate', 'named_executor_can_serve': True}
        assert determine_declinations(data) == []

    def test_intestate_spouse_pr(self):
        data = {'estate_type': 'Intestate', 'pr_relationship': 'Spouse'}
        assert determine_declinations(data) == []

    def test_intestate_child_pr_married_decedent(self):
        data = {
            'estate_type': 'Intestate',
            'pr_relationship': 'Child',
            'pr_full_name': 'Jane Doe',
            'decedent_marital_status': 'Married',
            'decedent_spouse_name': 'Mary Smith',
            'decedent_gender': 'Male',
            'heirs': [
                {'heir_full_name': 'Jane Doe', 'heir_relationship': 'Daughter',
                 'heir_gender': 'Female', 'heir_is_minor': False},
                {'heir_full_name': 'John Doe Jr', 'heir_relationship': 'Son',
                 'heir_gender': 'Male', 'heir_is_minor': False},
            ]
        }
        decls = determine_declinations(data)
        names = [d['name'] for d in decls]
        assert 'Mary Smith' in names      # Spouse must decline
        assert 'John Doe Jr' in names     # Sibling must decline
        assert 'Jane Doe' not in names    # PR doesn't decline

    def test_testate_executor_cant_serve_no_alternate(self):
        data = {
            'estate_type': 'Testate',
            'named_executor_can_serve': False,
            'will_executor_name': 'Old Executor',
            'will_executor_gender': 'Male',
            'will_names_alternate': False,
            'pr_full_name': 'Jane Doe',
            'heirs': [
                {'heir_full_name': 'Jane Doe', 'heir_relationship': 'Daughter',
                 'heir_is_beneficiary': True, 'heir_is_minor': False},
                {'heir_full_name': 'Bob Doe', 'heir_relationship': 'Son',
                 'heir_is_beneficiary': True, 'heir_is_minor': False},
            ]
        }
        decls = determine_declinations(data)
        names = [d['name'] for d in decls]
        assert 'Old Executor' in names   # Named executor must decline
        assert 'Bob Doe' in names        # Adult beneficiary must consent/decline
        assert 'Jane Doe' not in names   # PR excluded

    def test_testate_executor_cant_serve_with_alternate(self):
        data = {
            'estate_type': 'Testate',
            'named_executor_can_serve': False,
            'will_executor_name': 'Old Executor',
            'will_executor_gender': 'Male',
            'will_names_alternate': True,
            'will_alternate_name': 'Alternate Person',
        }
        decls = determine_declinations(data)
        names = [d['name'] for d in decls]
        assert 'Old Executor' in names
        assert len(decls) == 1  # Only named executor, no further declinations

    def test_intestate_grandchild_pr(self):
        data = {
            'estate_type': 'Intestate',
            'pr_relationship': 'Grandchild',
            'pr_full_name': 'Grandchild Doe',
            'decedent_marital_status': 'Single',
            'heirs': [
                {'heir_full_name': 'Child One', 'heir_relationship': 'Son',
                 'heir_gender': 'Male', 'heir_is_minor': False},
                {'heir_full_name': 'Child Two', 'heir_relationship': 'Daughter',
                 'heir_gender': 'Female', 'heir_is_minor': False},
            ]
        }
        decls = determine_declinations(data)
        names = [d['name'] for d in decls]
        assert 'Child One' in names
        assert 'Child Two' in names

    def test_intestate_minor_heirs_excluded(self):
        data = {
            'estate_type': 'Intestate',
            'pr_relationship': 'Child',
            'pr_full_name': 'Adult Child',
            'decedent_marital_status': 'Single',
            'heirs': [
                {'heir_full_name': 'Adult Child', 'heir_relationship': 'Son',
                 'heir_gender': 'Male', 'heir_is_minor': False},
                {'heir_full_name': 'Minor Child', 'heir_relationship': 'Son',
                 'heir_gender': 'Male', 'heir_is_minor': True},
            ]
        }
        decls = determine_declinations(data)
        names = [d['name'] for d in decls]
        assert 'Minor Child' not in names  # Minors can't decline


class TestSelectClosingDocuments:
    def test_testate_all_sui_juris(self):
        data = {'estate_type': 'Testate', 'all_heirs_sui_juris': True}
        docs = select_closing_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition to Close Estate CURLY.docx' in filenames
        assert 'Order to Close Estate CURLY.docx' in filenames

    def test_testate_not_all_sui_juris(self):
        data = {'estate_type': 'Testate', 'all_heirs_sui_juris': False}
        docs = select_closing_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition to Close Estate-No sui juris.docx' in filenames

    def test_intestate_closing(self):
        data = {'estate_type': 'Intestate'}
        docs = select_closing_documents(data)
        filenames = [d[0] for d in docs]
        assert 'Petition to Close Intestate Estate CURLY.docx' in filenames
        assert 'Order Closing Intestate Estate CURLY.docx' in filenames


class TestSelectReceiptWaiverTemplate:
    def test_intestate(self):
        heir = {'heir_full_name': 'Jane Doe'}
        data = {'estate_type': 'Intestate'}
        template, title = select_receipt_waiver_template(heir, data)
        assert template == 'Receipt and Waiver for Intestate Estate.docx'
        assert 'Jane Doe' in title

    def test_testate_residuary_and_pr(self):
        heir = {'heir_full_name': 'Jane Doe', 'heir_beneficiary_type': 'Residuary'}
        data = {'estate_type': 'Testate', 'pr_full_name': 'Jane Doe'}
        template, title = select_receipt_waiver_template(heir, data)
        assert template == 'Receipt and Waiver  - Residuary and Executor.docx'

    def test_testate_residuary_not_pr(self):
        heir = {'heir_full_name': 'Bob Smith', 'heir_beneficiary_type': 'Residuary'}
        data = {'estate_type': 'Testate', 'pr_full_name': 'Jane Doe'}
        template, title = select_receipt_waiver_template(heir, data)
        assert template == 'Receipt and Waiver  - Residuary CURLY.docx'

    def test_testate_specific(self):
        heir = {'heir_full_name': 'Bob Smith', 'heir_beneficiary_type': 'Specific'}
        data = {'estate_type': 'Testate', 'pr_full_name': 'Jane Doe'}
        template, title = select_receipt_waiver_template(heir, data)
        assert template == 'Receipt & Waiver - Testate CURLY.docx'

    def test_testate_general(self):
        heir = {'heir_full_name': 'Bob Smith', 'heir_beneficiary_type': 'General'}
        data = {'estate_type': 'Testate', 'pr_full_name': 'Jane Doe'}
        template, title = select_receipt_waiver_template(heir, data)
        assert template == 'Receipt and Waiver  - General CURLY.docx'


class TestCalculateDeadlines:
    def test_absolute_bar_date(self):
        result = calculate_deadlines('2026-04-10')
        assert result['absolute_bar_date'] == '2027-04-10'

    def test_with_publication_date(self):
        result = calculate_deadlines('2026-04-10', '2026-05-01')
        assert result['claims_deadline'] == '2026-09-01'
        assert 'exception_deadline' in result

    def test_exception_is_30_days_after_claims(self):
        result = calculate_deadlines('2026-04-10', '2026-05-01')
        from datetime import datetime, timedelta
        claims = datetime.strptime(result['claims_deadline'], '%Y-%m-%d')
        exception = datetime.strptime(result['exception_deadline'], '%Y-%m-%d')
        assert (exception - claims).days == 30

    def test_leap_year_dod(self):
        result = calculate_deadlines('2024-02-29')
        assert result['absolute_bar_date'] == '2025-02-28'

    def test_no_publication_date(self):
        result = calculate_deadlines('2026-04-10')
        assert 'claims_deadline' not in result
        assert 'exception_deadline' not in result


class TestGenerateFlags:
    def test_nonresident_pr(self):
        data = {'pr_state': 'Alabama'}
        flags = generate_flags(data)
        assert any('Nonresident' in f['message'] for f in flags)

    def test_small_estate_eligible(self):
        data = {'estimated_estate_value': 25000, 'small_estate_election': False}
        flags = generate_flags(data)
        assert any('small estate' in f['message'] for f in flags)

    def test_no_flags_for_clean_case(self):
        data = {'pr_state': 'Tennessee', 'decedent_county': 'Maury'}
        flags = generate_flags(data)
        assert len(flags) == 0

    def test_criminal_history(self):
        data = {'pr_criminal_history': True}
        flags = generate_flags(data)
        assert any('criminal history' in f['message'] for f in flags)

    def test_minor_heir_flag(self):
        data = {'heirs': [{'heir_is_minor': True}]}
        flags = generate_flags(data)
        assert any('sui juris' in f['message'] for f in flags)

    def test_out_of_county_property(self):
        data = {
            'decedent_county': 'Maury',
            'properties': [
                {'property_county': 'Maury'},
                {'property_county': 'Williamson'},
            ]
        }
        flags = generate_flags(data)
        assert any('Williamson' in f['message'] for f in flags)

    def test_business_interest(self):
        data = {'decedent_had_business': True}
        flags = generate_flags(data)
        assert any('business interest' in f['message'] for f in flags)

    def test_will_no_inventory_waiver(self):
        data = {'estate_type': 'Testate', 'will_waives_inventory': False}
        flags = generate_flags(data)
        assert any('inventory' in f['message'] for f in flags)


class TestBuildCommonReplacements:
    def test_builds_decedent_fields(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_county': 'Maury',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Female',
            'pr_full_name': 'Jane Smith',
        }
        result = build_common_replacements(data)
        assert result['{DECEDENT NAME}'] == 'John Smith'
        assert result['{DECEDENT}'] == 'John Smith'
        assert result['{HIS/HER}'] == 'his'
        assert result['{HE/SHE}'] == 'he'
        assert result['{DATE OF DEATH}'] == 'January 15, 2026'

    def test_builds_pr_fields(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Female',
            'pr_full_name': 'Jane Smith',
        }
        result = build_common_replacements(data)
        assert result['{PETITIONER NAME}'] == 'Jane Smith'
        assert result['{Executor/Executrix/PR Title CHOOSE ONE}'] == 'Executrix'
        assert result['{PETITIONER PRONOUNT – HE/SHE}'] == 'she'

    def test_criminal_history_clean(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
        }
        result = build_common_replacements(data)
        assert '{CRIMINAL_STATEMENT}' in result
        assert 'never been convicted' in result['{CRIMINAL_STATEMENT}']

    def test_criminal_history_disclosed(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'pr_criminal_history': True,
            'pr_criminal_details': 'convicted of DUI in 2010',
        }
        result = build_common_replacements(data)
        assert result['{CRIMINAL_STATEMENT}'] == 'convicted of DUI in 2010'

    def test_business_statement_no_business(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
        }
        result = build_common_replacements(data)
        assert '{BUSINESS_STATEMENT}' in result
        assert 'was not the owner' in result['{BUSINESS_STATEMENT}']

    def test_business_statement_with_business(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'decedent_had_business': True,
            'decedent_business_details': 'a sole proprietor of ABC Corp',
        }
        result = build_common_replacements(data)
        assert 'ABC Corp' in result['{BUSINESS_STATEMENT}']

    def test_waiver_statement(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'will_waives_bond': True,
            'will_waives_inventory': True,
        }
        result = build_common_replacements(data)
        assert '{WAIVER_STATEMENT}' in result
        assert 'making bond' in result['{WAIVER_STATEMENT}']
        assert 'filing an inventory' in result['{WAIVER_STATEMENT}']

    def test_waiver_statement_empty(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
        }
        result = build_common_replacements(data)
        assert result['{WAIVER_STATEMENT}'] == ''

    def test_sui_juris_statement_all_sui_juris(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'heirs': [
                {'heir_full_name': 'Jane Smith', 'heir_is_minor': False,
                 'heir_has_disability': False},
            ],
        }
        result = build_common_replacements(data)
        assert result['{SUI_JURIS_STATEMENT}'] == 'All beneficiaries are sui juris.'

    def test_sui_juris_statement_not_all(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'heirs': [
                {'heir_full_name': 'Minor Child', 'heir_is_minor': True},
            ],
        }
        result = build_common_replacements(data)
        assert result['{SUI_JURIS_STATEMENT}'] == 'Not all beneficiaries are sui juris.'

    def test_heirs_list_formatting(self):
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'heirs': [
                {'heir_full_name': 'Jane Smith', 'heir_age': 45,
                 'heir_relationship': 'Daughter', 'heir_address': '123 Main St',
                 'heir_city': 'Columbia'},
            ]
        }
        result = build_common_replacements(data)
        assert 'Jane Smith' in result['{HEIRS_LIST}']
        assert 'Daughter' in result['{HEIRS_LIST}']

    def test_placeholder_variants_petitioner(self):
        """Test that all petitioner name variants map correctly."""
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Female',
            'pr_full_name': 'Jane Smith',
            'pr_address': '123 Main St',
            'pr_city': 'Columbia',
            'pr_age': 45,
        }
        result = build_common_replacements(data)
        # All petitioner name variants
        assert result['{Petitioner Name}'] == 'Jane Smith'
        assert result['{petitioner}'] == 'Jane Smith'
        assert result["{petitioner's name}"] == 'Jane Smith'
        # PR address variants
        assert result['{ADDRESS OF PETITIONER}'] == '123 Main St'
        assert result['{STREET OF PETITIONER}'] == '123 Main St'
        assert result['{Petitioner Address}'] == '123 Main St'
        # PR city variants
        assert result['{CITY OF PETITIONER}'] == 'Columbia'
        # PR age variants
        assert result['{AGE OF PETITIONER}'] == '45'

    def test_placeholder_variants_decedent(self):
        """Test lowercase and mixed-case decedent variants."""
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'decedent_county': 'Maury',
            'decedent_age': 75,
            'estate_type': 'Testate',
            'pr_gender': 'Male',
        }
        result = build_common_replacements(data)
        assert result['{decedent name}'] == 'John Smith'
        assert result['{Date of Death}'] == 'January 15, 2026'
        assert result['{County of Residence for Decedent}'] == 'Maury'
        assert result['{AGE}'] == '75'
        assert result['{AGE AT DEATH}'] == '75'
        assert result['{DECEDENT PRONOUN \u2013 HIS/HER}'] == 'his'

    def test_placeholder_variants_attorney(self):
        """Test attorney name variants and BPR."""
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'attorney_full_name': 'Dale Hutto',
            'attorney_bpr': '12345',
        }
        result = build_common_replacements(data)
        assert result['{ATTORNEY}'] == 'Dale Hutto'
        assert result['{Attorney first name}'] == 'Dale'
        assert result['{BPR #}'] == '12345'

    def test_placeholder_variants_date(self):
        """Test year, month, and current year variants."""
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
        }
        result = build_common_replacements(data)
        from datetime import datetime
        now = datetime.now()
        assert result['{YEAR}'] == str(now.year)
        assert result['{Current Year}'] == str(now.year)
        assert result['{Month}'] == now.strftime('%B')

    def test_placeholder_variants_title_and_relationship(self):
        """Test title and relationship variants."""
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Female',
            'pr_full_name': 'Jane Smith',
            'pr_relationship': 'Daughter',
        }
        result = build_common_replacements(data)
        assert result['{Executor/Executrix}'] == 'Executrix'
        assert result['{Title}'] == 'Executrix'
        assert result['{Relation of Petition to Decedent}'] == 'Daughter'
        assert result['{Petitioner Pronoun HIS/HER}'] == 'her'

    def test_placeholder_docket_lowercase(self):
        """Test lowercase docket number variant."""
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'case_number': '2026-PR-001',
        }
        result = build_common_replacements(data)
        assert result['{docket number}'] == '2026-PR-001'

    def test_placeholder_were_or_no(self):
        """Test were/were no conditional placeholder."""
        data = {
            'decedent_full_name': 'John Smith',
            'decedent_gender': 'Male',
            'decedent_dod': '2026-01-15',
            'estate_type': 'Testate',
            'pr_gender': 'Male',
            'were_or_no_objections': 'were',
        }
        result = build_common_replacements(data)
        assert result['{were/were no}'] == 'were'


class TestDerivePRTitleErrors:
    def test_invalid_estate_type_raises_value_error(self):
        with pytest.raises(ValueError, match='Unknown estate_type/pr_gender'):
            derive_pr_title('InvalidType', 'Male')

    def test_invalid_gender_raises_value_error(self):
        with pytest.raises(ValueError, match='Unknown estate_type/pr_gender'):
            derive_pr_title('Testate', 'Other')

    def test_none_values_raise_value_error(self):
        with pytest.raises(ValueError):
            derive_pr_title(None, None)


class TestSelectOpeningDocumentsErrors:
    def test_testate_no_will_type_raises_value_error(self):
        data = {'estate_type': 'Testate'}
        with pytest.raises(ValueError, match='Unknown will_type'):
            select_opening_documents(data)

    def test_testate_invalid_will_type_raises_value_error(self):
        data = {'estate_type': 'Testate', 'will_type': 'Notarized'}
        with pytest.raises(ValueError, match='Unknown will_type'):
            select_opening_documents(data)
